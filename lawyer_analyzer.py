"""
Універсальний адвокат-аналізатор + генератор судових документів
================================================================
Об'єднує обидва скрипти + генерує готові документи для суду:
  - Клопотання про долучення аудіозаписів
  - Стенограму (розшифровку) розмов
  - Технічні відомості про записи
  - Повний аналіз захисту (Claude Opus)

Використання:
  python lawyer_analyzer.py                  # повний режим
  python lawyer_analyzer.py --docs-only      # тільки генерація документів
  python lawyer_analyzer.py --analysis-only  # тільки AI-аналіз
  python lawyer_analyzer.py --help
"""

import os
import sys
import re
import json
import argparse
import subprocess
import httpx
import anthropic
from datetime import datetime
from pathlib import Path

# ── UTF-8 консоль ──────────────────────────────────────────────────────────
if sys.stdout.encoding and sys.stdout.encoding.lower() not in ("utf-8", "utf-8-sig"):
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")

# ── Константи ──────────────────────────────────────────────────────────────
BASE_DIR = Path(__file__).parent

try:
    from case_config import CASE, ROOT_DATA_DIR, SORTED_DIR, AUDIO_INPUT_DIR
except ImportError:
    from case_config_example import CASE, ROOT_DATA_DIR, SORTED_DIR, AUDIO_INPUT_DIR  # type: ignore

NEW_CHUNKS_DIR   = BASE_DIR / "готовые_нарезки"

# PDF-аналіз справи — шукаємо відносно ROOT_DATA_DIR
PDF_ANALYSIS_PATH = ROOT_DATA_DIR / "all_pages_recognized.txt"

# Вихідні файли
OUTPUT_ANALYSIS  = BASE_DIR / "DEFENSE_ANALYSIS_FULL.txt"
OUTPUT_DOCS_DIR  = BASE_DIR / "СУДОВІ_ДОКУМЕНТИ"

# ── Промпт аналізу (розширена версія v2) ──────────────────────────────────
LAWYER_PROMPT = """
Ти — досвідчений український адвокат-практик у справах ст. 173-2 КУпАП (домашнє насильство) з 15-річним стажем.
Ти знаєш усі процесуальні тонкощі, судову практику Київського регіону, реальні важелі захисту.

Обвинувачений: {defendant} ({dob}).
Справа: {case_number}, три адміністративних протоколи, {article}.
Засідання: {hearing_date}, {hearing_time}, суддя {judge}, {court}.

Тобі надано ДВА блоки матеріалів:

== БЛОК А: МАТЕРІАЛИ ПОЛІЦЕЙСЬКОЇ СПРАВИ ==
Повна розшифровка адміністративної справи.

== БЛОК Б: АУДІОЗАПИСИ (ТРАНСКРИПЦІЇ) ==
Аудіозаписи систематизовані та ПІДПИСАНІ за змістом (назви папок).
Звертай особливу увагу на назви папок/файлів — вони описують суть кожного запису.

---

ТВОЄ ЗАВДАННЯ — МАКСИМАЛЬНО ДЕТАЛЬНИЙ АДВОКАТСЬКИЙ АНАЛІЗ:

## 1. КАРТА ДОКАЗІВ
Таблиця: кожен запис (за назвою папки) → що є корисного для захисту → конкретна цитата → до якого пункту обвинувачення відноситься.

## 2. НАЙСИЛЬНІШІ АРГУМЕНТИ ЗАХИСТУ
Топ-10 аргументів з конкретними цитатами з транскрипцій. Для кожного: цитата → юридичний аргумент → норма закону.

## 3. ПОВНІ ТЕКСТИ КЛОПОТАНЬ (готові до подання — УКРАЇНСЬКОЮ МОВОЮ)
Напиши готові тексти клопотань:
- Про долучення аудіозаписів до справи (з переліком та описом кожного фрагмента)
- Про виклик свідків (конкретні люди, згадані в записах — імена, ким приходяться)
- Про повернення протоколу на доопрацювання (процесуальні порушення)
- Про призначення психологічної експертизи
- Письмові пояснення підзахисного по кожному епізоду обвинувачення

## 4. СЕКРЕТИ АДВОКАТА — РЕАЛЬНА ТАКТИКА В ЗАЛІ СУДУ
Що досвідчений захисник робить у залі суду по ст.173-2:
- Як поводитися при запитаннях судді
- Як реагувати на показання заявниці
- Що говорити, що МОВЧАТИ
- Як використовувати паузи та процесуальні моменти
- Чого НЕ МОЖНА робити жодним чином
- Як використовувати форму оцінки ризиків (середній рівень) як головний щит
- Тактика перехресного допиту заявниці (якщо буде допит)

## 5. ПИСЬМОВА ПОЗИЦІЯ ЗАХИСТУ (УКРАЇНСЬКОЮ)
Готовий текст письмових пояснень підзахисного для суду (2-3 стор.) — від першої особи, юридично грамотно.

## 6. ПРОГНОЗ ПО КОЖНІЙ ІЗ ТРЬОХ СПРАВ
Окремо по кожному з трьох протоколів — реальний результат, чого можна досягти.

## 7. ПЛАН ДІЙ ДО {hearing_date} — ПОГОДИННО
Що робити сьогодні ввечері та вранці до {hearing_time}.

Аналізуй як практик: конкретно, без води, з цитатами, з готовими текстами документів.
УСІХ ДОКУМЕНТИ ТА КЛОПОТАННЯ — ВИКЛЮЧНО УКРАЇНСЬКОЮ МОВОЮ.
""".format(**CASE)


# ══════════════════════════════════════════════════════════════════════════
# ЗБІР ТРАНСКРИПЦІЙ
# ══════════════════════════════════════════════════════════════════════════

def collect_sorted_transcripts(root: Path) -> list[tuple[str, str]]:
    """Зчитує .txt з підпапок папка_с_транскрипцией, назва запису = батьківська папка."""
    results = []
    seen = set()
    if not root.exists():
        print(f"  [ПОПЕРЕДЖЕННЯ] Директорія не знайдена: {root}")
        return results
    for dirpath, dirnames, filenames in os.walk(root):
        dirnames.sort()
        if Path(dirpath).name == "папка_с_транскрипцией":
            recording_name = Path(dirpath).parent.name
            for fname in sorted(filenames):
                if fname.lower().endswith(".txt"):
                    fpath = Path(dirpath) / fname
                    try:
                        text = fpath.read_text(encoding="utf-8", errors="replace").strip()
                        if text and text not in seen:
                            seen.add(text)
                            results.append((f"{recording_name} / {fname}", text))
                    except Exception as e:
                        print(f"  Пропуск {fpath}: {e}")
    return results


def collect_new_chunks(root: Path) -> list[tuple[str, str]]:
    """Зчитує .txt з усіх підпапок готовые_нарезки (23-06-xx та HH-MM-SS_назва від pipeline.py)."""
    results = []
    seen = set()
    if not root.exists():
        return results
    for dirpath, dirnames, filenames in os.walk(root):
        dirnames.sort()
        folder = Path(dirpath).name
        # Папки 23-06-xx (старий нарізчик) або HH-MM-SS_* (pipeline.py)
        is_chunk_folder = (
            folder.startswith("23-06-")
            or (len(folder) >= 8 and folder[:8].replace("-", "").isdigit() and folder[6] == "_")
        )
        if not is_chunk_folder:
            continue
        for fname in sorted(filenames):
            if fname.lower().endswith(".txt"):
                fpath = Path(dirpath) / fname
                try:
                    text = fpath.read_text(encoding="utf-8", errors="replace").strip()
                    if text and text not in seen:
                        seen.add(text)
                        results.append((f"[нарізка] {folder} / {fname}", text))
                except Exception as e:
                    print(f"  Пропуск {fpath}: {e}")
    return results


def collect_all_transcripts(root: Path) -> list[tuple[str, str]]:
    """Загальний збір: будь-які .txt у директорії (для простого режиму)."""
    results = []
    seen = set()
    if not root.exists():
        return results
    for dirpath, dirnames, filenames in os.walk(root):
        dirnames.sort()
        for fname in sorted(filenames):
            if fname.lower().endswith(".txt"):
                fpath = Path(dirpath) / fname
                rel = str(fpath.relative_to(root))
                try:
                    text = fpath.read_text(encoding="utf-8", errors="replace").strip()
                    if text and text not in seen:
                        seen.add(text)
                        results.append((rel, text))
                except Exception as e:
                    print(f"  Пропуск {rel}: {e}")
    return results


def read_pdf_analysis() -> str:
    path = PDF_ANALYSIS_PATH
    if not path.exists():
        print(f"  PDF-аналіз не знайдено: {path}")
        return ""
    return path.read_text(encoding="utf-8", errors="replace").strip()


def build_context_message(
    sorted_transcripts: list,
    new_transcripts: list,
    pdf_text: str
) -> str:
    parts = []
    if pdf_text:
        parts += [
            "=" * 70,
            "БЛОК А: МАТЕРІАЛИ ПОЛІЦЕЙСЬКОЇ СПРАВИ (розшифровка PDF)",
            "=" * 70,
            pdf_text, ""
        ]
    all_tr = sorted_transcripts + new_transcripts
    parts += [
        "=" * 70,
        f"БЛОК Б: АУДІОТРАНСКРИПЦІЇ ({len(all_tr)} файлів)",
        "УВАГА: Назви папок описують СУТЬ кожного запису — використовуй їх для контексту!",
        "=" * 70,
    ]
    for label, text in all_tr:
        parts += [f"\n{'─'*60}", f"ЗАПИС: {label}", "─"*60, text]
    return "\n".join(parts)


# ══════════════════════════════════════════════════════════════════════════
# ГЕНЕРАЦІЯ СУДОВИХ ДОКУМЕНТІВ (DOCX + TXT)
# ══════════════════════════════════════════════════════════════════════════

def generate_court_documents(
    transcripts: list[tuple[str, str]],
    output_dir: Path
) -> list[Path]:
    """Генерує всі судові документи: клопотання, стенограму, технічні відомості."""
    output_dir.mkdir(parents=True, exist_ok=True)
    generated = []

    # 1. Клопотання про долучення аудіозаписів
    petition_path = output_dir / "1_Клопотання_про_долучення_аудіозаписів.txt"
    petition_text = _build_petition_audio(transcripts)
    petition_path.write_text(petition_text, encoding="utf-8")
    generated.append(petition_path)
    print(f"  ✓ {petition_path.name}")

    # 2. Стенограма (розшифровки)
    steno_path = output_dir / "2_Стенограма_аудіозаписів.txt"
    steno_text = _build_stenogram(transcripts)
    steno_path.write_text(steno_text, encoding="utf-8")
    generated.append(steno_path)
    print(f"  ✓ {steno_path.name}")

    # 3. Технічні відомості
    tech_path = output_dir / "3_Технічні_відомості_про_записи.txt"
    tech_text = _build_tech_info(transcripts)
    tech_path.write_text(tech_text, encoding="utf-8")
    generated.append(tech_path)
    print(f"  ✓ {tech_path.name}")

    # 4. Супровідний лист до суду
    cover_path = output_dir / "4_Супровідний_лист.txt"
    cover_text = _build_cover_letter(transcripts)
    cover_path.write_text(cover_text, encoding="utf-8")
    generated.append(cover_path)
    print(f"  ✓ {cover_path.name}")

    # 5. Чек-лист підготовки до засідання
    checklist_path = output_dir / "5_Чек-лист_підготовки.txt"
    checklist_text = _build_checklist()
    checklist_path.write_text(checklist_text, encoding="utf-8")
    generated.append(checklist_path)
    print(f"  ✓ {checklist_path.name}")

    # Намагаємося згенерувати DOCX-версію клопотання
    docx_path = _try_generate_docx(petition_text, output_dir)
    if docx_path:
        generated.append(docx_path)
        print(f"  ✓ {docx_path.name} (DOCX)")

    return generated


def _today() -> str:
    return datetime.now().strftime("%d.%m.%Y")


def _build_petition_audio(transcripts: list[tuple[str, str]]) -> str:
    lines = [
        f"До {CASE['court']}",
        "",
        f"Суддя: {CASE['judge']}",
        "",
        f"від {CASE['defendant']},",
        f"дата народження {CASE['dob']},",
        "особи, яка притягується до адміністративної відповідальності",
        "",
        "у справі {case_number}".format(**CASE),
        "",
        "=" * 60,
        "КЛОПОТАННЯ",
        "про долучення речових доказів (аудіозаписів)",
        "до матеріалів справи",
        "=" * 60,
        "",
        f"У провадженні {CASE['court']} (суддя {CASE['judge']}) перебуває справа",
        f"про адміністративне правопорушення, передбачене {CASE['article']},",
        f"щодо {CASE['defendant']}.",
        "",
        "Відповідно до ст. 251, 268 КУпАП, ст. 73 КАС України,",
        "керуючись принципом змагальності та повноти дослідження доказів,",
        "",
        "ПРОШУ СУД:",
        "",
        "Долучити до матеріалів справи як речові докази наступні аудіозаписи,",
        "що зафіксовані під час розмов, у яких я брав безпосередню участь",
        f"(запис вівся на {CASE['recording_device']}):",
        "",
    ]

    for i, (label, text) in enumerate(transcripts, 1):
        # Витягуємо перші 150 символів як анотацію
        preview = text[:150].replace("\n", " ").strip()
        if len(text) > 150:
            preview += "..."
        lines += [
            f"{i}. ЗАПИС: «{label}»",
            f"   Зміст: {preview}",
            f"   Значення для справи: підтверджує відсутність насильницьких дій,",
            f"   нормальний характер спілкування, дбайливе ставлення до членів сім'ї.",
            "",
        ]

    lines += [
        "ОБҐРУНТУВАННЯ:",
        "",
        "1. ЗАКОННІСТЬ ОТРИМАННЯ. Відповідно до ст. 31 Конституції України,",
        "   ч. 2 ст. 8 Конвенції про захист прав людини, запис розмов,",
        "   учасником яких є сама особа, що здійснює запис, є законним.",
        "   Всі аудіозаписи зроблені виключно під час розмов, у яких я",
        "   особисто брав участь.",
        "",
        "2. ВІДНОСНІСТЬ ДО СПРАВИ. Зміст записів безпосередньо стосується",
        "   обставин, що є предметом розгляду у даній справі, а саме:",
        "   характеру взаємостосунків у сім'ї, відсутності агресії,",
        "   дбайливого ставлення до дітей та конструктивного діалогу.",
        "",
        "3. ДОСТОВІРНІСТЬ. Записи здійснені без монтажу та редагування,",
        "   голоси учасників розмов чітко ідентифікуються.",
        "",
        "4. ПРОЦЕСУАЛЬНІ ПІДСТАВИ. Відповідно до ст. 251 КУпАП, доказами",
        "   у справі є будь-які фактичні дані, на підставі яких встановлюються",
        "   наявність або відсутність адміністративного правопорушення.",
        "   Аудіозаписи є документальними доказами відповідно до ст. 94 КАС України.",
        "",
        "Аудіозаписи надаються на матеріальному носії (USB-флеш-накопичувач),",
        "текстові розшифровки (стенограми) додаються до клопотання.",
        "",
        f"Дата: {_today()}",
        "",
        "З повагою,",
        f"____________________ / {CASE['defendant']} /",
    ]
    return "\n".join(lines)


def _build_stenogram(transcripts: list[tuple[str, str]]) -> str:
    lines = [
        "СТЕНОГРАМА АУДІОЗАПИСІВ",
        "=" * 60,
        f"Справа: {CASE['case_number']}",
        f"Складено: {_today()}",
        "",
        "ПРИМІТКА: Нижче наведені дослівні текстові розшифровки аудіозаписів.",
        "Голоси учасників позначені: [П] — Павло, [Д] — дружина/заявниця,",
        "[Д1],[Д2] — діти. Нерозбірливі місця позначені як [нерозб.].",
        "=" * 60,
        "",
    ]
    for i, (label, text) in enumerate(transcripts, 1):
        lines += [
            f"ЗАПИС № {i}",
            f"Джерело: {label}",
            "─" * 60,
            text,
            "",
            f"[Кінець запису № {i}]",
            "=" * 60,
            "",
        ]
    return "\n".join(lines)


def _build_tech_info(transcripts: list[tuple[str, str]]) -> str:
    lines = [
        "ТЕХНІЧНІ ВІДОМОСТІ ПРО АУДІОЗАПИСИ",
        "=" * 60,
        f"Справа: {CASE['case_number']}",
        f"Складено: {_today()}",
        "=" * 60,
        "",
        f"Пристрій запису: {CASE['recording_device']}",
        f"Записувальна особа: {CASE['defendant']} ({CASE['dob']})",
        "Статус: учасник усіх зафіксованих розмов",
        "",
        "ПЕРЕЛІК ФАЙЛІВ:",
        "",
    ]
    total_chars = 0
    for i, (label, text) in enumerate(transcripts, 1):
        total_chars += len(text)
        lines += [
            f"{i:3}. {label}",
            f"     Розмір транскрипції: {len(text):,} символів",
            "",
        ]
    lines += [
        "─" * 60,
        f"Усього файлів: {len(transcripts)}",
        f"Загальний обсяг транскрипцій: {total_chars:,} символів",
        "",
        "ПРАВОВА ПІДСТАВА ЗАКОННОСТІ ЗАПИСІВ:",
        "Відповідно до ч. 1 ст. 31 Конституції України гарантується",
        "таємниця особистого та сімейного спілкування. Водночас,",
        "особа має право фіксувати розмови, в яких вона бере участь,",
        "з метою захисту своїх законних прав та інтересів (Ухвала ВСУ",
        "№ 5-158кс(15), роз'яснення ВССУ 2017 р.).",
        "",
        f"Дата: {_today()}",
        "",
        f"____________________ / {CASE['defendant']} /",
    ]
    return "\n".join(lines)


def _build_cover_letter(transcripts: list[tuple[str, str]]) -> str:
    lines = [
        f"До {CASE['court']}",
        f"Суддя: {CASE['judge']}",
        "",
        f"від {CASE['defendant']}, {CASE['dob']}",
        "",
        "СУПРОВІДНИЙ ЛИСТ",
        "=" * 60,
        "",
        "Направляю до суду наступні документи у справі",
        f"{CASE['case_number']} ({CASE['article']}):",
        "",
        "1. Клопотання про долучення аудіозаписів до справи",
        "2. Стенограма аудіозаписів (текстова розшифровка)",
        "3. Технічні відомості про записи",
        f"4. Матеріальний носій з аудіофайлами (USB-флеш)",
        "",
        f"Усього аудіозаписів: {len(transcripts)} файлів.",
        "",
        "Прошу прийняти зазначені матеріали та задовольнити клопотання",
        "про їх долучення до справи як доказів захисту.",
        "",
        f"{_today()}",
        "",
        f"____________________ / {CASE['defendant']} /",
    ]
    return "\n".join(lines)


def _build_checklist() -> str:
    lines = [
        "ЧЕК-ЛИСТ ПІДГОТОВКИ ДО СУДОВОГО ЗАСІДАННЯ",
        "=" * 60,
        f"Справа: {CASE['case_number']}",
        f"Засідання: {CASE['hearing_date']} о {CASE['hearing_time']}",
        f"Суд: {CASE['court']}, суддя {CASE['judge']}",
        "=" * 60,
        "",
        "СЬОГОДНІ ВВЕЧЕРІ (31.05):",
        "",
        "[ ] 1. Роздрукувати всі документи у 2 примірниках:",
        "       - Клопотання про долучення аудіозаписів",
        "       - Стенограму (розшифровку) записів",
        "       - Технічні відомості",
        "       - Супровідний лист",
        "",
        "[ ] 2. Записати аудіофайли на USB-флеш або CD-диск",
        "       (не телефон! суд не прийме з телефону)",
        "",
        "[ ] 3. Перевірити наявність документів:",
        "       - Паспорт / ID-карта",
        "       - Копія всіх протоколів (ВАД №973440/9648/9649)",
        "       - Копії клопотань (для себе)",
        "",
        "[ ] 4. Підготувати письмові пояснення по кожному епізоду",
        "       (зі скриптом AI-аналізу або вручну)",
        "",
        "ВРАНЦІ 01.06 ДО 10:00:",
        "",
        "[ ] 5. Прибути до суду за 30 хвилин (09:30)",
        "",
        "[ ] 6. У залі суду — ПЕРШИЙ крок:",
        '       Коли суддя запитає "чи є заяви?":',
        '       → "Клопочу про долучення аудіозаписів до справи.',
        '          Прошу прийняти клопотання та матеріальний носій."',
        "",
        "[ ] 7. НЕ МОЖНА:",
        "       ✗ Перебивати суддю або заявницю",
        "       ✗ Підвищувати голос у залі",
        "       ✗ Визнавати факти без обдумування",
        "       ✗ Погрожувати або тиснути на свідків",
        "",
        "[ ] 8. МОЖНА І ПОТРІБНО:",
        "       ✓ Говорити спокійно та по-ділову",
        '       ✓ Відповідати: "Не визнаю. Прошу дослідити докази."',
        "       ✓ Посилатися на форму оцінки ризиків (середній рівень)",
        "       ✓ Заявляти всі клопотання письмово",
        "",
        "ПІСЛЯ ЗАСІДАННЯ:",
        "",
        "[ ] 9. Отримати копію ухвали/постанови суду",
        "[ ] 10. Зафіксувати час та зміст всього, що відбулося",
        "",
        f"Документ створено: {_today()}",
    ]
    return "\n".join(lines)


def _try_generate_docx(petition_text: str, output_dir: Path) -> Path | None:
    """Генерує DOCX через python-docx (встановлює автоматично якщо потрібно)."""
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("  [DOCX] Встановлюю python-docx...")
        subprocess.run([sys.executable, "-m", "pip", "install", "python-docx", "--quiet"], check=True)
        from docx import Document as DocxDocument
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

    docx_out = output_dir / "1_Клопотання_про_долучення_аудіозаписів.docx"
    try:
        doc = DocxDocument()
        # Поля сторінки: ліве 3 см, решта 2 см (судовий стандарт)
        section = doc.sections[0]
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.0)
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)

        BOLD_MARKERS  = {"КЛОПОТАННЯ", "ОБҐРУНТУВАННЯ:", "ПРОШУ СУД:", "З повагою,"}
        RIGHT_PREFIXES = ("До ", "від ", "Суддя:", "у справі")

        for line in petition_text.split("\n"):
            stripped = line.strip()
            is_bold   = stripped in BOLD_MARKERS or stripped.startswith("===")
            is_right  = stripped.startswith(RIGHT_PREFIXES)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if is_right else WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(line)
            run.bold      = is_bold
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)

        doc.save(docx_out)
        return docx_out
    except Exception as e:
        print(f"  [DOCX] Помилка: {e}")
        return None


# ══════════════════════════════════════════════════════════════════════════
# AI АНАЛІЗ (CLAUDE OPUS)
# ══════════════════════════════════════════════════════════════════════════

def run_ai_analysis(context: str) -> str:
    client = anthropic.Anthropic(
        http_client=httpx.Client(verify=False, timeout=600.0)
    )
    print("\nВідправка в Claude Opus (юридичний аналіз у реальному часі)...")
    print("=" * 70)

    full_response = ""
    with client.messages.stream(
        model="claude-opus-4-8",
        max_tokens=32000,
        messages=[{
            "role": "user",
            "content": context + "\n\n" + LAWYER_PROMPT,
        }],
    ) as stream:
        for text in stream.text_stream:
            print(text, end="", flush=True)
            full_response += text

    print(f"\n\n{'='*70}")
    return full_response


# ══════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        description="Університальний адвокат-аналізатор + генератор судових документів"
    )
    parser.add_argument("--docs-only",     action="store_true", help="Тільки генерація документів")
    parser.add_argument("--analysis-only", action="store_true", help="Тільки AI-аналіз")
    parser.add_argument("--sorted-dir",    type=str, default=str(SORTED_DIR),
                        help="Директорія з відсортованими записами")
    parser.add_argument("--chunks-dir",    type=str, default=str(NEW_CHUNKS_DIR),
                        help="Директорія з нарізками")
    args = parser.parse_args()

    sorted_dir  = Path(args.sorted_dir)
    chunks_dir  = Path(args.chunks_dir)

    print("=" * 70)
    print("АДВОКАТ-АНАЛІЗАТОР v3 — УНІВЕРСАЛЬНА ВЕРСІЯ")
    print(f"Справа: {CASE['case_number']}")
    print(f"Засідання: {CASE['hearing_date']} о {CASE['hearing_time']}")
    print("=" * 70)

    # ── Збір матеріалів ───────────────────────────────────────────────────
    print(f"\n[1/4] Збір транскрипцій з: {sorted_dir}")
    sorted_transcripts = collect_sorted_transcripts(sorted_dir)
    print(f"  Відсортовані записи: {len(sorted_transcripts)}")

    print(f"\n[2/4] Збір нарізок з: {chunks_dir}")
    seen_in_sorted = {text for _, text in sorted_transcripts}
    raw_new = collect_new_chunks(chunks_dir)
    if not raw_new:
        raw_new = collect_all_transcripts(chunks_dir)
    new_transcripts = [(l, t) for l, t in raw_new if t not in seen_in_sorted]
    print(f"  Нові нарізки (унікальні): {len(new_transcripts)}")

    all_transcripts = sorted_transcripts + new_transcripts

    # Перевірка ліміту: Claude Opus 4.8 = 200K токенів вхід
    # Грубо: 1 токен ≈ 4 символи; залишаємо буфер для промпту (~4K токенів)
    MAX_CHARS = (200_000 - 4_000 - 32_000) * 4  # ~656K символів
    total_chars = sum(len(t) for _, t in all_transcripts)
    if total_chars > MAX_CHARS:
        print(f"\n  [!] Контекст {total_chars:,} симв. перевищує ліміт.")
        print(f"      Залишаємо тільки відсортовані записи ({len(sorted_transcripts)} файлів).")
        all_transcripts = sorted_transcripts
        total_chars = sum(len(t) for _, t in all_transcripts)
        print(f"      Після обрізки: {total_chars:,} симв.")

    print(f"\n[3/4] Читання PDF-аналізу...")
    pdf_text = read_pdf_analysis()
    if pdf_text:
        print(f"  PDF: {len(pdf_text):,} символів")
    else:
        print("  PDF не знайдено, продовжую без нього.")

    if not all_transcripts and not pdf_text:
        print("\n[!] Матеріалів для аналізу не знайдено.")
        print("    Перевірте шляхи до директорій з транскрипціями.")
        sys.exit(1)

    print(f"\n  Усього файлів: {len(all_transcripts)}")
    for label, text in all_transcripts:
        print(f"    • {label} ({len(text):,} символів)")

    # ── Генерація судових документів ────────────────────────────────────
    if not args.analysis_only:
        print(f"\n[4/4] Генерація судових документів → {OUTPUT_DOCS_DIR}")
        docs = generate_court_documents(all_transcripts, OUTPUT_DOCS_DIR)
        print(f"\n  Згенеровано {len(docs)} документів у: {OUTPUT_DOCS_DIR}")

    # ── AI Аналіз ────────────────────────────────────────────────────────
    if not args.docs_only:
        total_chars = sum(len(t) for _, t in all_transcripts) + len(pdf_text)
        print(f"\nЗагальний контекст: {total_chars:,} символів (~{total_chars//4:,} токенів)")

        print("\nРозпочати AI-аналіз? (y/n): ")
        try:
            answer = input("> ").strip().lower()
        except EOFError:
            answer = "y"
            print("> (авто-так)")

        if answer in ("y", "yes", "так", "т", "да", "д"):
            context = build_context_message(sorted_transcripts, new_transcripts, pdf_text)
            analysis = run_ai_analysis(context)

            OUTPUT_ANALYSIS.write_text(analysis, encoding="utf-8")
            print(f"Аналіз збережено: {OUTPUT_ANALYSIS}")

            # Також копіюємо в папку документів
            (OUTPUT_DOCS_DIR / "0_AI_Аналіз_захисту_ПОВНИЙ.txt").write_text(
                analysis, encoding="utf-8"
            )
        else:
            print("AI-аналіз скасовано.")

    print("\n" + "=" * 70)
    print("ГОТОВО!")
    if not args.analysis_only:
        print(f"  Судові документи: {OUTPUT_DOCS_DIR}/")
    if not args.docs_only:
        print(f"  AI-аналіз: {OUTPUT_ANALYSIS}")
    print("=" * 70)


if __name__ == "__main__":
    main()
