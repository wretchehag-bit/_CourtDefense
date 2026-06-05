"""Service layer — per-job isolated pipeline.

Each task gets its own jobs/{tid}/ directory. No shared state between users.
Provides: transcription, PDF OCR, document extraction, classification, analysis.
"""
import os
import re
import uuid
import time
import json
from pathlib import Path
from threading import Lock, Thread
from typing import Dict, List, Optional, Tuple

from . import config

SCRIPT_DIR = config.PROJECT_ROOT
JOBS_DIR   = config.JOBS_DIR

AUDIO_EXTS = config.AUDIO_EXTENSIONS
DOC_EXTS   = config.DOC_EXTENSIONS

# ── Task storage ──────────────────────────────────────────────────────────
_tasks:   Dict[str, dict] = {}
_cancels: Dict[str, bool] = {}   # tid → True when user requests stop
_lock  = Lock()


CD_DIR = "_CourtDefense"   # subfolder created inside the user's source folder


def create_task(label: str = "") -> str:
    tid = uuid.uuid4().hex[:8]
    with _lock:
        _tasks[tid] = {
            "stage": "pending", "label": label or "Очікування",
            "progress": 0, "message": "", "logs": [], "files": [],
            "error": None, "current": None,
            "stage_files": {},   # stage → [{name, path, kb}]
            "source_dir": None,  # absolute path of the user's folder
        }
        _cancels[tid] = False
    return tid


def _add_stage_file(tid: str, stage: str, path: Path, base: Path):
    """Register a saved file under stage_files[stage]."""
    try:
        rel = str(path.relative_to(base)).replace("\\", "/")
    except ValueError:
        rel = str(path).replace("\\", "/")
    entry = {"name": path.name, "path": rel, "kb": round(path.stat().st_size / 1024, 1)}
    with _lock:
        if tid not in _tasks:
            return
        _tasks[tid]["stage_files"].setdefault(stage, []).append(entry)


def cancel_task(tid: str):
    with _lock:
        _cancels[tid] = True
        if tid in _tasks and _tasks[tid]["stage"] not in ("completed", "error"):
            _tasks[tid]["stage"] = "cancelling"
            _tasks[tid]["label"] = "Зупиняється…"


def _is_cancelled(tid: str) -> bool:
    return _cancels.get(tid, False)


def get_task(tid: str) -> Optional[dict]:
    with _lock:
        return dict(_tasks[tid]) if tid in _tasks else None


def all_tasks() -> dict:
    with _lock:
        return {k: dict(v) for k, v in _tasks.items()}


def _upd(tid: str, **kw):
    with _lock:
        if tid not in _tasks:
            return
        _tasks[tid].update(kw)
        msg = kw.get("message", "")
        if msg and msg.strip():
            _tasks[tid]["logs"].append(msg)
            _tasks[tid]["logs"] = _tasks[tid]["logs"][-200:]


# ── Audio transcription ───────────────────────────────────────────────────

def _run_transcriptions(
    tid: str,
    audio_files: List[Path],
    content_blocks: List[Tuple[str, str]],
    p_start: int, p_end: int,
    tr_dir: Optional[Path] = None,    # where to save .txt files
    base_dir: Optional[Path] = None,  # root for relative paths in stage_files
) -> bool:
    """Returns False if cancelled."""
    try:
        from faster_whisper import WhisperModel
    except ImportError:
        _upd(tid, message="[!] faster-whisper не встановлено: pip install faster-whisper")
        return False

    device, compute = "cpu", "int8"
    gpu_label = "CPU"
    try:
        import torch
        if torch.cuda.is_available():
            device, compute = "cuda", "float16"
            gpu_label = f"GPU: {torch.cuda.get_device_name(0)}"
    except ImportError:
        pass

    _upd(tid, message=f"Завантажую Whisper medium ({gpu_label})...")
    model = WhisperModel("medium", device=device, compute_type=compute)
    _upd(tid, message="Модель завантажена.", current=None)

    t_batch = time.time()

    for i, audio in enumerate(audio_files, 1):
        if _is_cancelled(tid):
            _upd(tid, stage="cancelled", label="Зупинено", progress=0,
                 message="Зупинено користувачем", current=None)
            return False

        pct = p_start + int((i - 1) / len(audio_files) * (p_end - p_start))
        t0 = time.time()
        _upd(tid, progress=pct,
             message=f"[{i}/{len(audio_files)}] Транскрибую: {audio.name}",
             current={"file": audio.name, "index": i,
                      "total": len(audio_files),
                      "elapsed": time.time() - t_batch,
                      "device": gpu_label})
        try:
            segments, _ = model.transcribe(
                str(audio), language="uk", beam_size=5,
                vad_filter=True,
                vad_parameters={"min_silence_duration_ms": 500},
            )
            lines = []
            for seg in segments:
                h = int(seg.start // 3600)
                m = int((seg.start % 3600) // 60)
                s = seg.start % 60
                lines.append(f"[{h:02d}:{m:02d}:{s:05.2f}] {seg.text.strip()}")
            text = "\n".join(lines)
            content_blocks.append((f"АУДІОЗАПИС: {audio.stem}", text))
            elapsed = time.time() - t0

            # Save transcript immediately to _CourtDefense/01_транскрипції/
            if tr_dir:
                tr_dir.mkdir(parents=True, exist_ok=True)
                tr_file = tr_dir / (audio.stem + ".txt")
                tr_file.write_text(text, encoding="utf-8")
                if base_dir:
                    _add_stage_file(tid, "transcribing", tr_file, base_dir)

            _upd(tid,
                 message=f"✓ {audio.name}: {len(lines)} сегм., {len(text):,} симв. ({elapsed:.0f}с)",
                 current={"file": audio.name, "index": i,
                          "total": len(audio_files),
                          "elapsed": time.time() - t_batch,
                          "device": gpu_label})
        except Exception as exc:
            _upd(tid, message=f"[!] Помилка транскрипції {audio.name}: {exc}")

    _upd(tid, current=None)
    return True


# ── PDF OCR via Claude Files API ──────────────────────────────────────────

_OCR_PROMPT = """Ти — фахівець з розпізнавання юридичних документів і рукописного тексту.
Розшифруй ВЕСЬ текст цього документу дослівно, сторінку за сторінкою.
Формат: для кожної сторінки пиши заголовок "=== СТОРІНКА N ===" перед текстом.
Якщо слово нерозбірливо — пиши [нерозбірливо].
Зберігай структуру: дати, ПІБ, адреси, підписи ([підпис]), штампи ([штамп: текст]).
Виводь тільки чистий текст без пояснень і коментарів."""

_OCR_BATCH_SIZE  = 50            # сторінок за один запит (< 100/хв rate limit)
_OCR_BATCH_DELAY = 65            # секунд між батчами
_FILE_SIZE_LIMIT = 30 * 1024 * 1024   # 30 MB per page-file


def _make_anthropic_client():
    import anthropic, httpx
    return anthropic.Anthropic(http_client=httpx.Client(verify=False, timeout=300.0))


def _upload_and_ask(client, page_files: list, prompt_suffix: str = "") -> str:
    """Upload page PDFs, ask Claude to transcribe, clean up."""
    uploaded = []
    try:
        for pf in page_files:
            with open(pf, "rb") as f:
                uploaded.append(client.beta.files.upload(
                    file=(pf.name, f, "application/pdf"),
                ))
        content = [
            {"type": "document", "source": {"type": "file", "file_id": u.id}}
            for u in uploaded
        ]
        content.append({"type": "text", "text": _OCR_PROMPT + prompt_suffix})

        result = ""
        with client.messages.stream(
            model="claude-opus-4-8",
            max_tokens=64000,
            messages=[{"role": "user", "content": content}],
            extra_headers={"anthropic-beta": "files-api-2025-04-14"},
        ) as stream:
            for chunk in stream.text_stream:
                result += chunk
        return result
    finally:
        for u in uploaded:
            try:
                client.beta.files.delete(u.id)
            except Exception:
                pass


def _ocr_pdf_single(path: Path, tid: str, client) -> str:
    """One-shot OCR for small PDFs (single file upload, ≤ 30 MB)."""
    _upd(tid, message=f"  Завантажую в Files API одним файлом…")
    uploaded = None
    try:
        with open(path, "rb") as f:
            uploaded = client.beta.files.upload(file=(path.name, f, "application/pdf"))
        _upd(tid, message=f"  Розпізнаю (id={uploaded.id[:12]}…)")
        result = ""
        with client.messages.stream(
            model="claude-opus-4-8",
            max_tokens=64000,
            messages=[{"role": "user", "content": [
                {"type": "document", "source": {"type": "file", "file_id": uploaded.id}},
                {"type": "text", "text": _OCR_PROMPT},
            ]}],
            extra_headers={"anthropic-beta": "files-api-2025-04-14"},
        ) as stream:
            for chunk in stream.text_stream:
                result += chunk
        return result
    except Exception as e:
        _upd(tid, message=f"  ❌ Помилка OCR: {str(e)}")
        raise
    finally:
        if uploaded:
            try:
                client.beta.files.delete(uploaded.id)
            except Exception:
                pass


def _ocr_pdf_chunked(path: Path, tid: str, client, n_pages: int) -> str:
    """Batch OCR: split PDF into page-files, process in chunks of _OCR_BATCH_SIZE."""
    import tempfile
    try:
        from pypdf import PdfReader, PdfWriter
    except ImportError:
        return "[pypdf не встановлено: pip install pypdf]"

    _upd(tid, message=f"  Розбиваю {n_pages} сторінок на окремі файли…")
    reader = PdfReader(str(path))

    with tempfile.TemporaryDirectory() as tmp:
        tmp_dir = Path(tmp)
        page_files = []
        for i, page in enumerate(reader.pages, 1):
            w = PdfWriter()
            w.add_page(page)
            out = tmp_dir / f"p{i:04d}.pdf"
            with open(out, "wb") as f:
                w.write(f)
            page_files.append(out)

        batches = [page_files[i:i + _OCR_BATCH_SIZE]
                   for i in range(0, len(page_files), _OCR_BATCH_SIZE)]
        n_batches = len(batches)
        all_parts = []

        for idx, batch in enumerate(batches, 1):
            if idx > 1:
                _upd(tid, message=f"  ⏳ Пауза {_OCR_BATCH_DELAY}с (rate limit)…")
                time.sleep(_OCR_BATCH_DELAY)

            p_start = (idx - 1) * _OCR_BATCH_SIZE + 1
            p_end   = p_start + len(batch) - 1
            suffix  = f"\n[Батч {idx}/{n_batches}: стор. {p_start}–{p_end}]" if n_batches > 1 else ""
            _upd(tid, message=f"  Батч {idx}/{n_batches}: стор. {p_start}–{p_end} — завантажую…")

            text = _upload_and_ask(client, batch, suffix)
            all_parts.append(text)
            _upd(tid, message=f"  ✓ Батч {idx}/{n_batches}: {len(text):,} символів")

    return "\n\n".join(all_parts)


def _ocr_pdf_auto(path: Path, tid: str) -> str:
    """Auto-select single vs chunked OCR strategy."""
    client  = _make_anthropic_client()
    size_mb = path.stat().st_size / 1024 / 1024

    # Count pages cheaply
    n_pages = 0
    try:
        from pypdf import PdfReader
        n_pages = len(PdfReader(str(path)).pages)
    except Exception:
        pass

    _upd(tid, message=f"  Сканований PDF: {n_pages or '?'} стор., {size_mb:.1f} MB")

    # Small PDF: single upload is fine (fast, 1 API call)
    if size_mb < 28 and n_pages <= 20:
        return _ocr_pdf_single(path, tid, client)

    # Large PDF: split & batch
    _upd(tid, message=f"  Великий PDF → пакетний OCR (~{max(1,(n_pages+_OCR_BATCH_SIZE-1)//_OCR_BATCH_SIZE)} батч(и))")
    return _ocr_pdf_chunked(path, tid, client, n_pages or 999)


# ── Standalone PDF-to-text pipeline ──────────────────────────────────────

def _pipeline_convert_pdf(tid: str):
    job_dir    = JOBS_DIR / tid
    upload_dir = job_dir / "uploads"
    output_dir = job_dir / "output"

    pdfs = sorted(f for f in upload_dir.glob("*") if f.suffix.lower() == ".pdf") \
           if upload_dir.exists() else []
    if not pdfs:
        _upd(tid, stage="error", error="PDF не знайдено", progress=100)
        return

    output_dir.mkdir(parents=True, exist_ok=True)
    files_out = []

    for i, pdf in enumerate(pdfs, 1):
        _upd(tid, stage="transcribing",
             label=f"OCR [{i}/{len(pdfs)}]: {pdf.name}",
             progress=int((i - 1) / len(pdfs) * 90) + 5,
             message=f"[{i}/{len(pdfs)}] Обробляю: {pdf.name}")

        # 1) Fast text extraction
        text_pages = []
        try:
            import pdfplumber
            with pdfplumber.open(str(pdf)) as doc:
                for j, page in enumerate(doc.pages, 1):
                    t = page.extract_text() or ""
                    if t.strip():
                        text_pages.append(f"=== СТОРІНКА {j} ===\n{t.strip()}")
        except Exception:
            pass

        if text_pages:
            _upd(tid, message=f"  Текстовий PDF — pdfplumber знайшов {len(text_pages)} стор.")
            full_text = "\n\n".join(text_pages)
        else:
            _upd(tid, message=f"  Сканований — запускаю OCR через Claude Files API…")
            try:
                full_text = _ocr_pdf_auto(pdf, tid)
            except Exception as e:
                error_msg = f"API помилка: {str(e)}"
                _upd(tid, message=f"  ❌ {error_msg}")
                full_text = f"[ПОМИЛКА OCR: {error_msg}]"

        out_txt = output_dir / (pdf.stem + "_text.txt")
        out_txt.write_text(full_text, encoding="utf-8")
        try:
            rel = str(out_txt.relative_to(SCRIPT_DIR)).replace("\\", "/")
        except ValueError:
            rel = str(out_txt).replace("\\", "/")
        files_out.append({"name": out_txt.name, "path": rel,
                          "kb": round(out_txt.stat().st_size / 1024, 1)})
        _upd(tid, message=f"  ✓ Збережено: {out_txt.name} ({files_out[-1]['kb']} KB)")

    _upd(tid, stage="completed", label="Конвертація завершена!", progress=100,
         files=files_out, message=f"✅ {len(pdfs)} PDF → текст")


def start_convert_pdf(tid: str):
    Thread(target=_pipeline_convert_pdf, args=(tid,), daemon=True).start()


# ── Document text extraction ──────────────────────────────────────────────

def _extract_doc(path: Path, tid: str = "") -> str:
    ext = path.suffix.lower()

    if ext == ".txt":
        try:
            return path.read_text(encoding="utf-8", errors="replace")
        except Exception as e:
            return f"[Помилка читання TXT: {e}]"

    if ext == ".docx":
        try:
            from docx import Document
            doc = Document(str(path))
            parts = []
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text)
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts)
        except ImportError:
            return "[python-docx не встановлено: pip install python-docx]"
        except Exception as e:
            return f"[Помилка DOCX: {e}]"

    if ext == ".pdf":
        # 1) Try fast text extraction first
        text_pages = []
        try:
            import pdfplumber
            with pdfplumber.open(str(path)) as pdf:
                for i, page in enumerate(pdf.pages, 1):
                    t = page.extract_text() or ""
                    if t.strip():
                        text_pages.append(f"--- стор. {i} ---\n{t.strip()}")
        except ImportError:
            pass
        except Exception:
            pass

        if text_pages:
            return "\n\n".join(text_pages)

        # 2) Scanned PDF → auto OCR (single or chunked depending on size)
        if tid:
            return _ocr_pdf_auto(path, tid)
        return "[PDF: текст не виявлено. Запусти через webapp для автоматичного OCR.]"

    return ""


def _run_extractions(
    tid: str,
    doc_files: List[Path],
    content_blocks: List[Tuple[str, str]],
    p_start: int, p_end: int,
):
    for i, doc in enumerate(doc_files, 1):
        pct = p_start + int((i - 1) / len(doc_files) * (p_end - p_start))
        _upd(tid, progress=pct, message=f"[{i}/{len(doc_files)}] Читаю: {doc.name}")
        try:
            text = _extract_doc(doc, tid)
            if text.strip():
                content_blocks.append((f"ДОКУМЕНТ: {doc.name}", text))
                _upd(tid, message=f"✓ {doc.name}: {len(text):,} символів")
            else:
                _upd(tid, message=f"[!] {doc.name}: порожній вміст")
        except Exception as exc:
            _upd(tid, message=f"[!] Помилка {doc.name}: {exc}")


def _run_extractions_with_save(
    tid: str,
    doc_files: List[Path],
    content_blocks: List[Tuple[str, str]],
    doc_dir: Path,
    base_dir: Path,
    p_start: int, p_end: int,
):
    """Extract docs and save extracted text to doc_dir."""
    for i, doc in enumerate(doc_files, 1):
        pct = p_start + int((i - 1) / len(doc_files) * (p_end - p_start))
        _upd(tid, progress=pct, message=f"[{i}/{len(doc_files)}] Читаю: {doc.name}")
        try:
            text = _extract_doc(doc, tid)
            if text.strip():
                content_blocks.append((f"ДОКУМЕНТ: {doc.name}", text))
                saved = doc_dir / (doc.stem + "_text.txt")
                saved.write_text(text, encoding="utf-8")
                _add_stage_file(tid, "extracting", saved, base_dir)
                _upd(tid, message=f"✓ {doc.name}: {len(text):,} символів")
            else:
                _upd(tid, message=f"[!] {doc.name}: порожній вміст")
        except Exception as exc:
            _upd(tid, message=f"[!] Помилка {doc.name}: {exc}")


# ── Shared helpers ────────────────────────────────────────────────────────

def _get_case() -> dict:
    """Load CASE from case_config.py or return empty dict."""
    try:
        from case_config import CASE
        return CASE
    except ImportError:
        try:
            from case_config_example import CASE  # type: ignore
            return CASE
        except ImportError:
            return {}


def _checkpoint(path: Path, min_size: int = 100) -> bool:
    """True if file exists, is readable and has enough content."""
    return path.exists() and path.is_file() and path.stat().st_size >= min_size


def _make_client():
    import anthropic, httpx
    return anthropic.Anthropic(http_client=httpx.Client(verify=False, timeout=300.0))


def _read_txt_safe(path: Path) -> str:
    for enc in ("utf-8", "utf-8-sig", "cp1251", "latin1"):
        try:
            return path.read_text(encoding=enc, errors="strict")
        except (UnicodeDecodeError, LookupError):
            continue
    return path.read_text(encoding="utf-8", errors="replace")


# ── Stage 2 — Classify & Organize (advocate agent logic) ─────────────────

_CLASSIFY_SYSTEM = """\
Ти — досвідчений адвокат по справах КУпАП з 15-річним стажем у судах України.
Оціни аудіотранскрипцію як ДОКАЗ ЗАХИСТУ.
{case_ctx}
STRONG  — заявник суперечить своїм показанням; спокійний тон обвинуваченого; турбота про дітей; провокація з боку заявника.
SUPPORT — нейтральне спілкування без агресії; побутові питання вирішуються мирно.
NEUTRAL — технічні розмови, короткі репліки не пов'язані зі справою.
RISKY   — підвищений голос; лексика що трактується як тиск/погрози; обвинувачений виглядає агресором.
Мова матеріалів може бути мішаниною — аналізуй усі мови.
Відповідай ТІЛЬКИ JSON без markdown.\
"""

_CLASSIFY_USER = """\
Назва запису: "{name}"

Транскрипція (до 5 000 симв.):
{text}

{{
  "category": "STRONG|SUPPORT|NEUTRAL|RISKY",
  "score": <ціле 1-10>,
  "reason": "<1-2 речення>",
  "key_quotes": ["<цитата>"],
  "court_tip": "<1 речення>"
}}\
"""

_CATEGORY_DIRS = {
    "STRONG":  "1_СИЛЬНІ_ДОКАЗИ",
    "SUPPORT": "2_ДОПОМІЖНІ",
    "NEUTRAL": "3_НЕЙТРАЛЬНІ",
    "RISKY":   "4_РИЗИКОВІ_НЕ_НЕСТИ",
}


def _classify_one(client, recording_name: str, text: str, case_ctx: str) -> dict:
    """Classify a single transcript. Returns dict with category/score/etc."""
    snippet = text[:5000] + ("\n...[скорочено]" if len(text) > 5000 else "")
    user_msg = _CLASSIFY_USER.format(name=recording_name, text=snippet)
    system   = _CLASSIFY_SYSTEM.format(case_ctx=case_ctx)

    for attempt in range(3):
        try:
            resp = client.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=600,
                system=system,
                messages=[{"role": "user", "content": user_msg}],
            )
            raw  = resp.content[0].text.strip()
            raw  = re.sub(r"^```(?:json)?\s*|\s*```$", "", raw, flags=re.MULTILINE)
            data = json.loads(raw)
            cat  = data.get("category", "NEUTRAL")
            if cat not in _CATEGORY_DIRS:
                cat = "NEUTRAL"
            return {
                "category":   cat,
                "score":      max(1, min(10, int(data.get("score", 5)))),
                "reason":     data.get("reason", ""),
                "key_quotes": data.get("key_quotes", [])[:3],
                "court_tip":  data.get("court_tip", ""),
            }
        except json.JSONDecodeError:
            if attempt < 2:
                time.sleep(2 ** attempt)
        except Exception:
            if attempt < 2:
                time.sleep(3)

    return {"category": "NEUTRAL", "score": 5, "reason": "Помилка класифікації",
            "key_quotes": [], "court_tip": "Перевірити вручну"}


def _run_classify_organize(
    tid: str,
    tr_dir: Path,           # 01_транскрипції/
    doc_texts: List[Tuple[str, str]],  # [(label, text)] extracted docs
    org_dir: Path,          # 02_організовано/
    base_dir: Path,
    api_key: str,
    case: dict,
) -> List[Tuple[str, str]]:
    """
    Classify every transcript with Claude (or just copy if no key).
    Returns content_blocks = [(label, text)] for the analysis step.
    """
    import json as _json

    cache_path = org_dir / ".classify_cache.json"
    cache: dict = {}
    if cache_path.exists():
        try:
            cache = _json.loads(cache_path.read_text(encoding="utf-8"))
        except Exception:
            cache = {}

    # Find all transcripts
    txt_files = sorted(tr_dir.glob("*.txt")) if tr_dir.exists() else []
    if not txt_files:
        _upd(tid, message="  [!] Транскрипцій не знайдено для організації")
        return list(doc_texts)

    for cat_dir in _CATEGORY_DIRS.values():
        (org_dir / cat_dir).mkdir(parents=True, exist_ok=True)

    case_ctx = ""
    if case.get("defendant") and "Прізвище" not in case.get("defendant", ""):
        case_ctx = (f"\nСправа: {case.get('case_number','')}, {case.get('article','')}"
                    f"\nОбвинувачений: {case.get('defendant','')}\n")

    client = _make_client() if api_key else None
    content_blocks: List[Tuple[str, str]] = list(doc_texts)

    for i, txt_path in enumerate(txt_files, 1):
        recording_name = txt_path.stem
        pct_inner = int((i - 1) / len(txt_files) * 100)

        # Check cache (checkpoint)
        cache_key = recording_name
        if cache_key in cache:
            _upd(tid, message=f"  ⚡ [{i}/{len(txt_files)}] кеш: {recording_name[:55]}")
            result = cache[cache_key]
        elif not api_key:
            # No key: put all in NEUTRAL, skip classification
            result = {"category": "NEUTRAL", "score": 5, "reason": "Без API ключа",
                      "key_quotes": [], "court_tip": ""}
        else:
            _upd(tid, message=f"  🔍 [{i}/{len(txt_files)}] класифікую: {recording_name[:50]}")
            text   = _read_txt_safe(txt_path)
            result = _classify_one(client, recording_name, text, case_ctx)
            cache[cache_key] = result
            try:
                cache_path.write_text(
                    json.dumps(cache, ensure_ascii=False, indent=2), encoding="utf-8")
            except Exception:
                pass

        cat = result["category"]
        icons = {"STRONG": "✅", "SUPPORT": "⚠️", "NEUTRAL": "🔍", "RISKY": "❌"}
        _upd(tid, message=f"    {icons.get(cat,'?')} {recording_name[:55]} → {cat} (score:{result['score']})")
        _add_stage_file(tid, "extracting",
                        txt_path, base_dir)  # original transcript file

        # Copy to category folder
        safe  = re.sub(r'[<>:"/\\|?*]', '_', recording_name)
        dest  = org_dir / _CATEGORY_DIRS[cat] / f"score{result['score']:02d}_{safe}.txt"
        if not dest.exists():
            import shutil
            shutil.copy2(txt_path, dest)

        # Add to context blocks (STRONG and SUPPORT only — keep context lean)
        if cat in ("STRONG", "SUPPORT"):
            text = _read_txt_safe(txt_path)
            block_label = f"АУДІО [{cat}] {recording_name}"
            content_blocks.append((block_label, text))

    # Generate CLASSIFICATION_REPORT.txt
    report_lines = ["ЗВІТ КЛАСИФІКАЦІЇ АУДІОЗАПИСІВ", "="*60, ""]
    by_cat: dict = {c: [] for c in _CATEGORY_DIRS}
    for name, res in cache.items():
        by_cat.get(res.get("category", "NEUTRAL"), by_cat["NEUTRAL"]).append(
            (name, res.get("score", 5), res.get("reason", ""))
        )
    for cat, dirname in _CATEGORY_DIRS.items():
        items = sorted(by_cat[cat], key=lambda x: -x[1])
        report_lines.append(f"\n{dirname} ({len(items)} файл(ів)):")
        for name, score, reason in items:
            report_lines.append(f"  [{score:2}] {name[:70]}")
            if reason:
                report_lines.append(f"       {reason[:100]}")

    report = "\n".join(report_lines)
    report_file = org_dir / "КЛАСИФІКАЦІЯ.txt"
    report_file.write_text(report, encoding="utf-8")
    _add_stage_file(tid, "extracting", report_file, base_dir)
    _upd(tid, message=f"  📄 Збережено: КЛАСИФІКАЦІЯ.txt")

    return content_blocks


# ── Stage 3 — Defense Master analysis ────────────────────────────────────

_ANALYSIS_SYSTEM = """\
Ти — провідний адвокат-практик по адміністративних справах з 15-річним досвідом у судах України.
{case_ctx}
Аналізуй матеріали справи конкретно, з посиланнями на документи, без загальних фраз.
Мова матеріалів — суміш українська + російська, аналізуй обидві.\
"""

_ANALYSIS_USER = """\
Нижче всі матеріали справи. Підготуй повний пакет захисту.

{docs_block}

Структура відповіді:

=== 1. ЗАГАЛЬНА КАРТИНА СПРАВИ ===
Хто, що, коли. Кількість протоколів. Суть обвинувачення.

=== 2. СИЛЬНІ ПОЗИЦІЇ ЗАХИСТУ ===
Конкретні факти і цитати що допомагають обвинуваченому.

=== 3. СЛАБКІ МІСЦЯ / РИЗИКИ ===
Що може бути використано проти. Будь чесним.

=== 4. СУПЕРЕЧНОСТІ В МАТЕРІАЛАХ ===
Де обвинувачення суперечить само собі або доказам захисту.

=== 5. ПРОЦЕСУАЛЬНІ ПОРУШЕННЯ ===
Якщо є — конкретні норми і що порушено.

=== 6. СТРАТЕГІЯ ЗАХИСТУ ===
Конкретний план: що говорити, що подавати, порядок дій.

=== 7. КЛЮЧОВІ АРГУМЕНТИ ДО СУДУ ===
По пунктах, готові формулювання для виголошення.

=== 8. ПИТАННЯ ДО СВІДКІВ / ЗАЯВНИЦІ ===
Конкретні питання що підривають позицію обвинувачення.

=== 9. ШПАРГАЛКА ДО СУДУ ===
1 сторінка найважливішого — читати прямо в залі.\
"""


def _run_defense_master(
    tid: str,
    content_blocks: List[Tuple[str, str]],
    anal_dir: Path,
    base_dir: Path,
    case: dict,
) -> List[Path]:
    """Send all collected content to Claude and save final analysis files."""
    import anthropic, httpx

    anal_dir.mkdir(parents=True, exist_ok=True)

    # Checkpoint: analysis already done?
    analysis_file = anal_dir / "00_АНАЛІЗ_СПРАВИ.txt"
    if _checkpoint(analysis_file, min_size=500):
        _upd(tid, message=f"⚡ Аналіз вже є: {analysis_file.name} — беремо готовий")
        saved = [analysis_file]
        cheat_file = anal_dir / "00_ШПАРГАЛКА_В_ЗАЛ_СУДУ.txt"
        if cheat_file.exists():
            saved.append(cheat_file)
        return saved

    # Build documents block (respect context limit)
    MAX_TOTAL = 180_000
    parts, total_chars = [], 0
    for label, text in content_blocks:
        chunk = f"{'='*60}\n{label}\n{'='*60}\n{text}"
        if total_chars + len(chunk) > MAX_TOTAL:
            remaining = MAX_TOTAL - total_chars
            if remaining > 200:
                chunk = chunk[:remaining] + "\n...[скорочено]"
                parts.append(chunk)
            break
        parts.append(chunk)
        total_chars += len(chunk)
    docs_block = "\n\n".join(parts)

    case_ctx = ""
    if case.get("defendant") and "Прізвище" not in case.get("defendant", ""):
        case_ctx = (f"\nСправа: {case.get('case_number','')}, {case.get('article','')}"
                    f"\nОбвинувачений: {case.get('defendant','')} ({case.get('dob','')})"
                    f"\nЗасідання: {case.get('hearing_date','')} о {case.get('hearing_time','')}"
                    f", суддя {case.get('judge','')}, {case.get('court','')}\n")

    system = _ANALYSIS_SYSTEM.format(case_ctx=case_ctx)
    user   = _ANALYSIS_USER.format(docs_block=docs_block)

    _upd(tid, message=f"Надсилаю {total_chars:,} символів в Claude ({len(content_blocks)} блок(ів))…")

    client = _make_client()
    result = ""
    with client.messages.stream(
        model="claude-sonnet-4-6",
        max_tokens=16000,
        system=system,
        messages=[{"role": "user", "content": user}],
    ) as stream:
        for chunk in stream.text_stream:
            result += chunk

    _upd(tid, message=f"Claude відповів: {len(result):,} символів")

    # Save full analysis
    analysis_file.write_text(result, encoding="utf-8")
    saved = [analysis_file]
    _upd(tid, message=f"📄 Збережено: {analysis_file.name}")

    # Extract cheat sheet (section 9)
    m = re.search(r"={3,}\s*9[\.\s]+ШПАРГАЛКА[^\n]*\n(.*?)(?:={3,}|$)",
                  result, re.DOTALL | re.IGNORECASE)
    cheat = m.group(1).strip() if m else result[-3000:]
    if cheat:
        cheat_file = anal_dir / "00_ШПАРГАЛКА_В_ЗАЛ_СУДУ.txt"
        cheat_file.write_text(cheat, encoding="utf-8")
        saved.append(cheat_file)
        _upd(tid, message=f"📄 Збережено: {cheat_file.name}")

    return saved


# ── Old analyze_case kept for compatibility (used by convert-pdf pipeline) ─

def _analyze_case(
    tid: str,
    content_blocks: List[Tuple[str, str]],
    output_dir: Path,
) -> List[Path]:
    """Thin wrapper — delegates to _run_defense_master."""
    case = _get_case()
    return _run_defense_master(tid, content_blocks, output_dir, output_dir.parent, case)


# ── Claude analysis ───────────────────────────────────────────────────────

def _build_system_prompt(case: dict) -> str:
    if case.get("defendant") and case["defendant"] not in ("Прізвище Ім'я По-батькові", ""):
        case_block = f"""
ДАНІ СПРАВИ:
• Обвинувачений: {case.get('defendant', '—')}  ({case.get('dob', '—')})
• Справа: {case.get('case_number', '—')},  {case.get('article', '—')}
• Засідання: {case.get('hearing_date', '—')} о {case.get('hearing_time', '—')}
• Суддя: {case.get('judge', '—')},  {case.get('court', '—')}
"""
    else:
        case_block = ""

    return f"""Ти — досвідчений адвокат захисту з 15-річним стажем у судах України.
{case_block}
Тобі надані матеріали справи (аудіотранскрипції та/або документи).
Мова матеріалів — суміш українська + російська, аналізуй обидві.

Підготуй структурований пакет захисту у такому форматі:

=== 1. КЛАСИФІКАЦІЯ МАТЕРІАЛІВ ===
Для кожного аудіо/документа: СИЛЬНИЙ / ДОПОМІЖНИЙ / НЕЙТРАЛЬНИЙ / РИЗИКОВИЙ
(+ коротка причина + ключова цитата якщо є)

=== 2. ЗАГАЛЬНА КАРТИНА СПРАВИ ===
Що відбулось, хто сторони, суть обвинувачення.

=== 3. СИЛЬНІ ПОЗИЦІЇ ЗАХИСТУ ===
Конкретні факти, цитати, докази що допомагають обвинуваченому.

=== 4. СЛАБКІ МІСЦЯ / РИЗИКИ ===
Що може бути використано проти — будь чесним.

=== 5. СТРАТЕГІЯ ЗАХИСТУ ===
Конкретний план: що говорити, що подавати, якого порядку дотримуватись.

=== 6. КЛЮЧОВІ АРГУМЕНТИ (для виголошення в залі) ===
По пунктах, готові формулювання.

=== 7. КЛОПОТАННЯ ===
Які подавати, коли, щоб проситити.

=== 8. ПИТАННЯ ДО СВІДКІВ / ЗАЯВНИЦІ ===
Конкретні питання що підривають позицію обвинувачення.

=== 9. ШПАРГАЛКА ДО СУДУ ===
Одна сторінка найважливішого — для читання прямо в залі."""


def _analyze_case(
    tid: str,
    content_blocks: List[Tuple[str, str]],
    output_dir: Path,
) -> List[Path]:
    import anthropic, httpx

    try:
        from case_config import CASE
    except ImportError:
        try:
            from case_config_example import CASE  # type: ignore
        except ImportError:
            CASE = {}

    # Build combined text
    parts = []
    for label, text in content_blocks:
        parts.append(f"{'='*60}\n{label}\n{'='*60}\n{text}")
    full_text = "\n\n".join(parts)

    # Truncate to stay well within Claude's context
    if len(full_text) > 190_000:
        full_text = full_text[:190_000] + "\n\n[... скорочено через обсяг ...]"

    system_prompt = _build_system_prompt(CASE)

    _upd(tid, message=f"Надсилаю {len(full_text):,} символів у Claude...")

    client = anthropic.Anthropic(
        http_client=httpx.Client(verify=False, timeout=300.0)
    )

    result = ""
    with client.messages.stream(
        model="claude-sonnet-4-6",
        max_tokens=16000,
        system=system_prompt,
        messages=[{"role": "user", "content": f"МАТЕРІАЛИ ДЛЯ АНАЛІЗУ:\n\n{full_text}"}],
    ) as stream:
        for chunk in stream.text_stream:
            result += chunk

    _upd(tid, message=f"Claude відповів: {len(result):,} символів")

    output_dir.mkdir(parents=True, exist_ok=True)
    saved = []

    analysis_file = output_dir / "00_АНАЛІЗ_СПРАВИ.txt"
    analysis_file.write_text(result, encoding="utf-8")
    saved.append(analysis_file)
    _upd(tid, message=f"📄 Збережено: {analysis_file}")

    # Extract cheat sheet
    m = re.search(r"={3,}\s*9[\.\s]+ШПАРГАЛКА[^\n]*\n(.*?)(?:={3,}|$)", result, re.DOTALL | re.IGNORECASE)
    cheat = m.group(1).strip() if m else ""
    if not cheat and len(result) > 1000:
        cheat = result[-3000:]
    if cheat:
        cheat_file = output_dir / "00_ШПАРГАЛКА_В_ЗАЛ_СУДУ.txt"
        cheat_file.write_text(cheat, encoding="utf-8")
        saved.append(cheat_file)
        _upd(tid, message=f"Збережено: {cheat_file.name}")

    return saved


# ── Transcript resume helpers ─────────────────────────────────────────────

def _safe_name(stem: str) -> str:
    """Strip trailing dots/spaces that Windows forbids in dir names."""
    return re.sub(r'[\s.]+$', '', stem).strip()


def _find_sorted_dir(source_dir: Path) -> Optional[Path]:
    """
    Search for Отсортированные_данные_для_суда in:
    - source dir itself
    - parent dir
    - all immediate subdirs of parent (sibling folders)
    - script dir
    Returns first match, or None.
    """
    target = "Отсортированные_данные_для_суда"

    # Direct and obvious locations first
    for base in [source_dir, source_dir.parent, SCRIPT_DIR]:
        c = base / target
        if c.exists() and c.is_dir():
            return c

    # Sibling folders of source_dir (e.g. нарезки/Отсортированные)
    try:
        for sibling in sorted(source_dir.parent.iterdir()):
            if sibling.is_dir() and sibling != source_dir:
                c = sibling / target
                if c.exists() and c.is_dir():
                    return c
    except PermissionError:
        pass

    return None


def _transcript_path(audio: Path, sorted_dir: Path) -> Path:
    """Expected transcript path for a given audio file."""
    return sorted_dir / _safe_name(audio.stem) / "папка_с_транскрипцией" / "part_01_chunk.txt"


def _has_transcript(audio: Path, sorted_dir: Optional[Path] = None) -> bool:
    # 1) New format: _CourtDefense/01_транскрипції/{stem}.txt  (in same dir as audio)
    cd_txt = audio.parent / CD_DIR / "01_транскрипції" / (audio.stem + ".txt")
    if cd_txt.exists() and cd_txt.stat().st_size > 100:
        return True
    # 2) Old format: Отсортированные_данные_для_суда/{name}/папка_с_транскрипцией/part_01_chunk.txt
    if sorted_dir:
        t = _transcript_path(audio, sorted_dir)
        if t.exists() and t.stat().st_size > 100:
            return True
    return False


def _read_existing_transcript(audio: Path, sorted_dir: Optional[Path] = None) -> str:
    """Read transcript from disk — new format first, then old."""
    cd_txt = audio.parent / CD_DIR / "01_транскрипції" / (audio.stem + ".txt")
    if cd_txt.exists() and cd_txt.stat().st_size > 100:
        return cd_txt.read_text(encoding="utf-8", errors="replace")
    if sorted_dir:
        t = _transcript_path(audio, sorted_dir)
        if t.exists() and t.stat().st_size > 100:
            return t.read_text(encoding="utf-8", errors="replace")
    return ""


def check_audio_transcripts(audio_files: List[Path], sorted_dir: Optional[Path]) -> dict:
    """
    Split audio_files into already-done and needs-processing.
    Returns dict with lists + counts.
    """
    if not sorted_dir:
        return {
            "sorted_dir": None,
            "done":  [],
            "todo":  audio_files,
            "done_count": 0,
            "todo_count": len(audio_files),
        }

    done, todo = [], []
    for audio in audio_files:
        if _has_transcript(audio, sorted_dir):
            done.append(audio)
        else:
            todo.append(audio)


    return {
        "sorted_dir":  str(sorted_dir),
        "done":        done,
        "todo":        todo,
        "done_count":  len(done),
        "todo_count":  len(todo),
    }


# ── Folder scanner ────────────────────────────────────────────────────────

_SKIP_DIRS = {
    ".venv", "venv", ".env", "__pycache__", "node_modules",
    ".git", "build", "dist", "build_tmp", "jobs",
    "ПАКЕТ_ЗАХИСТУ", "ВІДІБРАНІ_ДОКАЗИ", "готовые_нарезки",
    CD_DIR,   # skip our own output folder when scanning
}


def scan_folder(folder_path: str) -> dict:
    """Recursively find all supported files, skipping service/build directories."""
    root = Path(folder_path)
    if not root.exists() or not root.is_dir():
        return {"error": f"Папка не знайдена: {folder_path}", "audio": [], "docs": []}

    audio, docs = [], []

    def _walk(directory: Path):
        try:
            entries = sorted(directory.iterdir())
        except PermissionError:
            return
        for entry in entries:
            if entry.is_dir():
                if entry.name in _SKIP_DIRS or entry.name.startswith("."):
                    continue
                _walk(entry)
            elif entry.is_file():
                ext = entry.suffix.lower()
                if ext in AUDIO_EXTS:
                    audio.append(entry)
                elif ext in DOC_EXTS:
                    docs.append(entry)

    _walk(root)
    audio.sort()
    docs.sort()

    return {
        "audio": audio,
        "docs":  docs,
        "total": len(audio) + len(docs),
        "error": None,
    }


# ── Pipeline core (shared by upload and folder modes) ─────────────────────

def _run_pipeline_core(tid: str, audio_files: List[Path],
                       doc_files: List[Path], source_dir: Path,
                       sorted_dir: Optional[Path] = None,
                       skip_existing: bool = True):
    """
    Unified pipeline:
      Step 1 (transcribing) — Whisper per audio file
      Step 2 (extracting)   — Classify transcripts (advocate agent) + extract docs
      Step 3 (analyzing)    — Defense master analysis (Claude)
    All output → source_dir/_CourtDefense/
    """
    cd_root  = source_dir / CD_DIR
    tr_dir   = cd_root / "01_транскрипції"
    org_dir  = cd_root / "02_організовано"
    anal_dir = cd_root / "03_аналіз"
    cd_root.mkdir(parents=True, exist_ok=True)

    case    = _get_case()
    api_key = os.environ.get("ANTHROPIC_API_KEY", "").strip()

    _upd(tid, source_dir=str(source_dir))

    content_blocks: List[Tuple[str, str]] = []

    # ── Step 1: Transcribe audio (5 → 45%) ───────────────────────────────
    if audio_files:
        if skip_existing:
            check = check_audio_transcripts(audio_files, sorted_dir)
            done_files = check["done"]
            todo_files = check["todo"]
        else:
            done_files, todo_files = [], audio_files

        # Load already-done from disk
        if done_files:
            _upd(tid, message=f"⚡ Вже є транскрипції: {len(done_files)} — читаю з диску")
            for audio in done_files:
                text = _read_existing_transcript(audio, sorted_dir)
                if text:
                    content_blocks.append((f"АУДІОЗАПИС: {audio.stem}", text))
                    _upd(tid, message=f"  ✓ (кеш) {audio.name}: {len(text):,} симв.")

        # Transcribe new files
        if todo_files:
            _upd(tid, stage="transcribing",
                 label=f"Транскрипція ({len(todo_files)} файл(ів))", progress=5,
                 message=f"Нових файлів: {len(todo_files)}")
            ok = _run_transcriptions(tid, todo_files, content_blocks,
                                     p_start=5, p_end=45,
                                     tr_dir=tr_dir, base_dir=source_dir)
            if not ok:
                return

    if _is_cancelled(tid):
        return

    # ── Step 2: Classify + Organize (40 → 65%) ───────────────────────────
    # a) Extract text from uploaded docs
    doc_texts: List[Tuple[str, str]] = []
    if doc_files:
        _upd(tid, stage="extracting",
             label=f"Читання документів ({len(doc_files)})", progress=40,
             message=f"Документів: {len(doc_files)}")
        for i, doc in enumerate(doc_files, 1):
            _upd(tid, progress=40 + int(i / len(doc_files) * 5),
                 message=f"  [{i}/{len(doc_files)}] Читаю: {doc.name}")
            text = _extract_doc(doc, tid)
            if text.strip():
                saved_doc = org_dir / "документи" / (doc.stem + "_text.txt")
                saved_doc.parent.mkdir(parents=True, exist_ok=True)
                saved_doc.write_text(text, encoding="utf-8")
                _add_stage_file(tid, "extracting", saved_doc, source_dir)
                doc_texts.append((f"ДОКУМЕНТ: {doc.name}", text))

    # b) Classify transcripts with advocate-agent logic
    _upd(tid, stage="extracting",
         label="Класифікація транскрипцій", progress=45,
         message=f"Класифікую {len(list(tr_dir.glob('*.txt')) if tr_dir.exists() else [])} транскрипцій…")

    content_blocks = _run_classify_organize(
        tid, tr_dir, doc_texts, org_dir, source_dir, api_key, case)

    if _is_cancelled(tid):
        return

    # ── Collect output files helper ───────────────────────────────────────
    def _collect_files():
        out = []
        if cd_root.exists():
            for f in sorted(cd_root.rglob("*.txt")):
                out.append({"name": f.name,
                             "path": str(f).replace("\\", "/"),
                             "kb":   round(f.stat().st_size / 1024, 1)})
        return out

    if not content_blocks:
        files_out = _collect_files()
        if files_out:
            _upd(tid, stage="completed", label="Транскрипції збережено",
                 progress=100, files=files_out,
                 message=f"⚠️ Нема контенту для аналізу, але транскрипції збережено.\n📁 {tr_dir}")
        else:
            _upd(tid, stage="error", progress=100,
                 error="Не вдалося отримати текст з жодного файлу",
                 message="[!] Перевір формат файлів або встанови залежності")
        return

    total_chars = sum(len(t) for _, t in content_blocks)
    _upd(tid, message=f"Зібрано {len(content_blocks)} блок(ів), {total_chars:,} символів")

    if _is_cancelled(tid):
        return

    # ── Step 3: Defense master analysis (65 → 99%) ───────────────────────
    if not api_key:
        files_out = _collect_files()
        tr_count  = sum(1 for f in files_out if "01_транскрипції" in f["path"])
        _upd(tid, stage="completed",
             label=f"✅ Транскрипції збережено ({tr_count} файл(ів))",
             progress=100, files=files_out,
             message=f"✅ Транскрипції: {tr_dir}\n"
                     f"   Класифікація: {org_dir}\n"
                     f"⚠️ Аналіз Claude пропущено — API ключ не вказано.\n"
                     f"   Введи ключ і запусти знову — Whisper не повториться.")
        return

    _upd(tid, stage="analyzing", label="Аналіз справи (Defense Master)", progress=65,
         message="Готую фінальний аналіз…")
    try:
        saved = _run_defense_master(tid, content_blocks, anal_dir, source_dir, case)
        for f in saved:
            _add_stage_file(tid, "analyzing", f, source_dir)
    except Exception as exc:
        files_out = _collect_files()
        _upd(tid, stage="error", error=str(exc), progress=100, files=files_out,
             message=f"[!] Claude помилка: {exc}\n"
                     f"✅ Транскрипції і класифікація збережені в {cd_root}\n"
                     f"   Виправ помилку і запусти знову — Whisper не повториться.")
        return

    # ── Done ──────────────────────────────────────────────────────────────
    files_out = _collect_files()
    _upd(tid, stage="completed", label="Готово!", progress=100,
         files=files_out, message=f"✅ Готово! {cd_root}")


# ── Pipeline: from uploaded files ─────────────────────────────────────────

def _pipeline(tid: str):
    upload_dir = JOBS_DIR / tid / "uploads"

    all_files   = sorted(upload_dir.glob("*")) if upload_dir.exists() else []
    audio_files = [f for f in all_files if f.suffix.lower() in AUDIO_EXTS]
    doc_files   = [f for f in all_files if f.suffix.lower() in DOC_EXTS]

    if not audio_files and not doc_files:
        _upd(tid, stage="error", error="Немає файлів для обробки", progress=100,
             message="[!] Завантаж аудіо або документи (.pdf, .docx, .txt)")
        return

    # For upload mode: save to upload_dir/_CourtDefense/
    _run_pipeline_core(tid, audio_files, doc_files, source_dir=upload_dir,
                       skip_existing=False)


def start_pipeline(tid: str):
    Thread(target=_pipeline, args=(tid,), daemon=True).start()


# ── Pipeline: from local folder (no upload) ───────────────────────────────

def _pipeline_from_folder(tid: str, folder_path: str, skip_existing: bool = True):
    result = scan_folder(folder_path)
    if result["error"]:
        _upd(tid, stage="error", error=result["error"], progress=100)
        return

    audio_files = result["audio"]
    doc_files   = result["docs"]

    if not audio_files and not doc_files:
        _upd(tid, stage="error", progress=100,
             error=f"Підтримуваних файлів не знайдено у: {folder_path}",
             message="[!] Шукаємо: " + ", ".join(sorted(AUDIO_EXTS | DOC_EXTS)))
        return

    source_dir = Path(folder_path)
    sorted_dir = _find_sorted_dir(source_dir)

    msg = f"Знайдено: {len(audio_files)} аудіо + {len(doc_files)} документів"
    if skip_existing and audio_files:
        check = check_audio_transcripts(audio_files, sorted_dir)
        msg += f" | ⚡ {check['done_count']} вже є, 🔄 {check['todo_count']} нових"
    _upd(tid, stage="uploaded", progress=5, message=msg)

    _run_pipeline_core(tid, audio_files, doc_files, source_dir=source_dir,
                       sorted_dir=sorted_dir, skip_existing=skip_existing)


def start_folder_pipeline(tid: str, folder_path: str, skip_existing: bool = True):
    Thread(target=_pipeline_from_folder,
           args=(tid, folder_path, skip_existing), daemon=True).start()


# ── Audio Cutter Pipeline ─────────────────────────────────────────────────
def start_audio_cutting(tid: str, folder_path: str, phrases_file: str):
    """Запустити нарізку аудіо в окремому потоці."""
    Thread(target=_audio_cutting_worker,
           args=(tid, folder_path, phrases_file), daemon=True).start()


def _audio_cutting_worker(tid: str, folder_path: str, phrases_file: str):
    """Worker thread для нарізки аудіо за фразами."""
    try:
        from .audio_cutter import cut_audio_by_timestamps

        folder = Path(folder_path)
        phrases_path = Path(phrases_file)

        _upd(tid, stage="processing", progress=10,
             message=f"🔍 Сканую папку та парсю фрази...")

        # Читаємо файл з фразами
        phrases_text = _read_text_safe(phrases_path, fallback_text="")
        if not phrases_text.strip():
            _upd(tid, stage="error", progress=100,
                 error=f"Файл фраз порожній або не прочитаний: {phrases_path}")
            return

        _upd(tid, stage="processing", progress=20,
             message=f"📂 Запускаю pipeline нарізки з {len(phrases_text.splitlines())} фраз...")

        # Запускаємо основну функцію нарізки
        result = cut_audio_by_timestamps(str(folder))

        if not result:
            _upd(tid, stage="error", progress=100,
                 error="Нарізка не повернула результат")
            return

        stats = result.get("stats", {})
        success_count = stats.get("success", 0)
        error_count = stats.get("errors", 0)
        total = success_count + error_count

        if success_count > 0:
            _upd(tid, stage="completed", progress=100,
                 message=f"✅ Готово! Нарізано {success_count} фрагментів" +
                        (f" ({error_count} помилок)" if error_count > 0 else ""))
        else:
            _upd(tid, stage="error", progress=100,
                 error=f"Помилка обробки. Перевірте файли та права доступу.")

    except Exception as e:
        import traceback
        err_msg = f"{str(e)}\n\nДеталі:\n{traceback.format_exc()}"
        _upd(tid, stage="error", progress=100, error=err_msg)


def list_results() -> list:
    results = []
    if not JOBS_DIR.exists():
        return results
    for job in sorted(JOBS_DIR.iterdir(), reverse=True):
        out = job / "output"
        if out.exists():
            for f in sorted(out.rglob("*.txt")):
                rel = str(f.relative_to(SCRIPT_DIR)).replace("\\", "/")
                results.append({"name": f.name, "path": rel,
                                 "kb": round(f.stat().st_size / 1024, 1)})
    return results


def resolve_file(rel_path: str) -> Optional[Path]:
    p = (SCRIPT_DIR / rel_path.replace("/", os.sep)).resolve()
    try:
        p.relative_to(SCRIPT_DIR.resolve())
        return p if p.exists() else None
    except ValueError:
        return None
