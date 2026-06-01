"""
Service layer — per-job isolated pipeline.
Each task gets its own jobs/{tid}/ directory. No shared state between users.
"""
import os, sys, re, uuid, time
from pathlib import Path
from threading import Lock, Thread
from typing import Dict, List, Optional, Tuple

SCRIPT_DIR = Path(__file__).parent.parent
JOBS_DIR   = SCRIPT_DIR / "jobs"

AUDIO_EXTS = {".mp3", ".m4a", ".wav", ".flac", ".ogg", ".aac", ".wma"}
DOC_EXTS   = {".pdf", ".docx", ".txt"}

# ── Task storage ──────────────────────────────────────────────────────────
_tasks: Dict[str, dict] = {}
_lock  = Lock()


def create_task(label: str = "") -> str:
    tid = uuid.uuid4().hex[:8]
    with _lock:
        _tasks[tid] = {
            "stage": "pending", "label": label or "Очікування",
            "progress": 0, "message": "", "logs": [], "files": [], "error": None,
        }
    return tid


def get_task(tid: str) -> Optional[dict]:
    with _lock:
        return dict(_tasks[tid]) if tid in _tasks else None


def all_tasks() -> dict:
    with _lock:
        return {k: dict(v) for k, v in _tasks.items()}


def _upd(tid: str, **kw):
    with _lock:
        if tid not in _tasks:
            return
        _tasks[tid].update(kw)
        msg = kw.get("message", "")
        if msg and msg.strip():
            _tasks[tid]["logs"].append(msg)
            _tasks[tid]["logs"] = _tasks[tid]["logs"][-200:]


# ── Audio transcription ───────────────────────────────────────────────────

def _run_transcriptions(
    tid: str,
    audio_files: List[Path],
    content_blocks: List[Tuple[str, str]],
    p_start: int, p_end: int,
):
    try:
        from faster_whisper import WhisperModel
    except ImportError:
        _upd(tid, message="[!] faster-whisper не встановлено: pip install faster-whisper")
        return

    device, compute = "cpu", "int8"
    try:
        import torch
        if torch.cuda.is_available():
            device, compute = "cuda", "float16"
    except ImportError:
        pass

    _upd(tid, message=f"Завантажую Whisper medium ({device.upper()})...")
    model = WhisperModel("medium", device=device, compute_type=compute)
    _upd(tid, message="Модель завантажена.")

    for i, audio in enumerate(audio_files, 1):
        pct = p_start + int((i - 1) / len(audio_files) * (p_end - p_start))
        _upd(tid, progress=pct, message=f"[{i}/{len(audio_files)}] Транскрибую: {audio.name}")
        try:
            segments, _ = model.transcribe(
                str(audio), language="uk", beam_size=5,
                vad_filter=True,
                vad_parameters={"min_silence_duration_ms": 500},
            )
            lines = []
            for seg in segments:
                h = int(seg.start // 3600)
                m = int((seg.start % 3600) // 60)
                s = seg.start % 60
                lines.append(f"[{h:02d}:{m:02d}:{s:05.2f}] {seg.text.strip()}")
            text = "\n".join(lines)
            content_blocks.append((f"АУДІОЗАПИС: {audio.stem}", text))
            _upd(tid, message=f"✓ {audio.name}: {len(lines)} сегментів, {len(text):,} символів")
        except Exception as exc:
            _upd(tid, message=f"[!] Помилка транскрипції {audio.name}: {exc}")


# ── PDF OCR via Claude Files API ──────────────────────────────────────────

_OCR_PROMPT = """Ти — фахівець з розпізнавання юридичних документів і рукописного тексту.
Розшифруй ВЕСЬ текст цього документу дослівно, сторінку за сторінкою.
Формат: для кожної сторінки пиши заголовок "=== СТОРІНКА N ===" перед текстом.
Якщо слово нерозбірливо — пиши [нерозбірливо].
Зберігай структуру: дати, ПІБ, адреси, підписи ([підпис]), штампи ([штамп: текст]).
Виводь тільки чистий текст без пояснень і коментарів."""

_OCR_BATCH_SIZE  = 50            # сторінок за один запит (< 100/хв rate limit)
_OCR_BATCH_DELAY = 65            # секунд між батчами
_FILE_SIZE_LIMIT = 30 * 1024 * 1024   # 30 MB per page-file


def _make_anthropic_client():
    import anthropic, httpx
    return anthropic.Anthropic(http_client=httpx.Client(verify=False, timeout=300.0))


def _upload_and_ask(client, page_files: list, prompt_suffix: str = "") -> str:
    """Upload page PDFs, ask Claude to transcribe, clean up."""
    uploaded = []
    try:
        for pf in page_files:
            with open(pf, "rb") as f:
                uploaded.append(client.beta.files.upload(
                    file=(pf.name, f, "application/pdf"),
                ))
        content = [
            {"type": "document", "source": {"type": "file", "file_id": u.id}}
            for u in uploaded
        ]
        content.append({"type": "text", "text": _OCR_PROMPT + prompt_suffix})

        result = ""
        with client.messages.stream(
            model="claude-opus-4-8",
            max_tokens=64000,
            messages=[{"role": "user", "content": content}],
            extra_headers={"anthropic-beta": "files-api-2025-04-14"},
        ) as stream:
            for chunk in stream.text_stream:
                result += chunk
        return result
    finally:
        for u in uploaded:
            try:
                client.beta.files.delete(u.id)
            except Exception:
                pass


def _ocr_pdf_single(path: Path, tid: str, client) -> str:
    """One-shot OCR for small PDFs (single file upload, ≤ 30 MB)."""
    _upd(tid, message=f"  Завантажую в Files API одним файлом…")
    with open(path, "rb") as f:
        uploaded = client.beta.files.upload(file=(path.name, f, "application/pdf"))
    _upd(tid, message=f"  Розпізнаю (id={uploaded.id[:12]}…)")
    result = ""
    try:
        with client.messages.stream(
            model="claude-opus-4-8",
            max_tokens=64000,
            messages=[{"role": "user", "content": [
                {"type": "document", "source": {"type": "file", "file_id": uploaded.id}},
                {"type": "text", "text": _OCR_PROMPT},
            ]}],
            extra_headers={"anthropic-beta": "files-api-2025-04-14"},
        ) as stream:
            for chunk in stream.text_stream:
                result += chunk
    finally:
        try:
            client.beta.files.delete(uploaded.id)
        except Exception:
            pass
    return result


def _ocr_pdf_chunked(path: Path, tid: str, client, n_pages: int) -> str:
    """Batch OCR: split PDF into page-files, process in chunks of _OCR_BATCH_SIZE."""
    import tempfile
    try:
        from pypdf import PdfReader, PdfWriter
    except ImportError:
        return "[pypdf не встановлено: pip install pypdf]"

    _upd(tid, message=f"  Розбиваю {n_pages} сторінок на окремі файли…")
    reader = PdfReader(str(path))

    with tempfile.TemporaryDirectory() as tmp:
        tmp_dir = Path(tmp)
        page_files = []
        for i, page in enumerate(reader.pages, 1):
            w = PdfWriter()
            w.add_page(page)
            out = tmp_dir / f"p{i:04d}.pdf"
            with open(out, "wb") as f:
                w.write(f)
            page_files.append(out)

        batches = [page_files[i:i + _OCR_BATCH_SIZE]
                   for i in range(0, len(page_files), _OCR_BATCH_SIZE)]
        n_batches = len(batches)
        all_parts = []

        for idx, batch in enumerate(batches, 1):
            if idx > 1:
                _upd(tid, message=f"  ⏳ Пауза {_OCR_BATCH_DELAY}с (rate limit)…")
                time.sleep(_OCR_BATCH_DELAY)

            p_start = (idx - 1) * _OCR_BATCH_SIZE + 1
            p_end   = p_start + len(batch) - 1
            suffix  = f"\n[Батч {idx}/{n_batches}: стор. {p_start}–{p_end}]" if n_batches > 1 else ""
            _upd(tid, message=f"  Батч {idx}/{n_batches}: стор. {p_start}–{p_end} — завантажую…")

            text = _upload_and_ask(client, batch, suffix)
            all_parts.append(text)
            _upd(tid, message=f"  ✓ Батч {idx}/{n_batches}: {len(text):,} символів")

    return "\n\n".join(all_parts)


def _ocr_pdf_auto(path: Path, tid: str) -> str:
    """Auto-select single vs chunked OCR strategy."""
    client  = _make_anthropic_client()
    size_mb = path.stat().st_size / 1024 / 1024

    # Count pages cheaply
    n_pages = 0
    try:
        from pypdf import PdfReader
        n_pages = len(PdfReader(str(path)).pages)
    except Exception:
        pass

    _upd(tid, message=f"  Сканований PDF: {n_pages or '?'} стор., {size_mb:.1f} MB")

    # Small PDF: single upload is fine (fast, 1 API call)
    if size_mb < 28 and n_pages <= 20:
        return _ocr_pdf_single(path, tid, client)

    # Large PDF: split & batch
    _upd(tid, message=f"  Великий PDF → пакетний OCR (~{max(1,(n_pages+_OCR_BATCH_SIZE-1)//_OCR_BATCH_SIZE)} батч(и))")
    return _ocr_pdf_chunked(path, tid, client, n_pages or 999)


# ── Standalone PDF-to-text pipeline ──────────────────────────────────────

def _pipeline_convert_pdf(tid: str):
    job_dir    = JOBS_DIR / tid
    upload_dir = job_dir / "uploads"
    output_dir = job_dir / "output"

    pdfs = sorted(f for f in upload_dir.glob("*") if f.suffix.lower() == ".pdf") \
           if upload_dir.exists() else []
    if not pdfs:
        _upd(tid, stage="error", error="PDF не знайдено", progress=100)
        return

    output_dir.mkdir(parents=True, exist_ok=True)
    files_out = []

    for i, pdf in enumerate(pdfs, 1):
        _upd(tid, stage="transcribing",
             label=f"OCR [{i}/{len(pdfs)}]: {pdf.name}",
             progress=int((i - 1) / len(pdfs) * 90) + 5,
             message=f"[{i}/{len(pdfs)}] Обробляю: {pdf.name}")

        # 1) Fast text extraction
        text_pages = []
        try:
            import pdfplumber
            with pdfplumber.open(str(pdf)) as doc:
                for j, page in enumerate(doc.pages, 1):
                    t = page.extract_text() or ""
                    if t.strip():
                        text_pages.append(f"=== СТОРІНКА {j} ===\n{t.strip()}")
        except Exception:
            pass

        if text_pages:
            _upd(tid, message=f"  Текстовий PDF — pdfplumber знайшов {len(text_pages)} стор.")
            full_text = "\n\n".join(text_pages)
        else:
            _upd(tid, message=f"  Сканований — запускаю OCR через Claude Files API…")
            full_text = _ocr_pdf_auto(pdf, tid)

        out_txt = output_dir / (pdf.stem + "_text.txt")
        out_txt.write_text(full_text, encoding="utf-8")
        rel = str(out_txt.relative_to(SCRIPT_DIR)).replace("\\", "/")
        files_out.append({"name": out_txt.name, "path": rel,
                          "kb": round(out_txt.stat().st_size / 1024, 1)})
        _upd(tid, message=f"  ✓ Збережено: {out_txt.name} ({files_out[-1]['kb']} KB)")

    _upd(tid, stage="completed", label="Конвертація завершена!", progress=100,
         files=files_out, message=f"✅ {len(pdfs)} PDF → текст")


def start_convert_pdf(tid: str):
    Thread(target=_pipeline_convert_pdf, args=(tid,), daemon=True).start()


# ── Document text extraction ──────────────────────────────────────────────

def _extract_doc(path: Path, tid: str = "") -> str:
    ext = path.suffix.lower()

    if ext == ".txt":
        try:
            return path.read_text(encoding="utf-8", errors="replace")
        except Exception as e:
            return f"[Помилка читання TXT: {e}]"

    if ext == ".docx":
        try:
            from docx import Document
            doc = Document(str(path))
            parts = []
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text)
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts)
        except ImportError:
            return "[python-docx не встановлено: pip install python-docx]"
        except Exception as e:
            return f"[Помилка DOCX: {e}]"

    if ext == ".pdf":
        # 1) Try fast text extraction first
        text_pages = []
        try:
            import pdfplumber
            with pdfplumber.open(str(path)) as pdf:
                for i, page in enumerate(pdf.pages, 1):
                    t = page.extract_text() or ""
                    if t.strip():
                        text_pages.append(f"--- стор. {i} ---\n{t.strip()}")
        except ImportError:
            pass
        except Exception:
            pass

        if text_pages:
            return "\n\n".join(text_pages)

        # 2) Scanned PDF → auto OCR (single or chunked depending on size)
        if tid:
            return _ocr_pdf_auto(path, tid)
        return "[PDF: текст не виявлено. Запусти через webapp для автоматичного OCR.]"

    return ""


def _run_extractions(
    tid: str,
    doc_files: List[Path],
    content_blocks: List[Tuple[str, str]],
    p_start: int, p_end: int,
):
    for i, doc in enumerate(doc_files, 1):
        pct = p_start + int((i - 1) / len(doc_files) * (p_end - p_start))
        _upd(tid, progress=pct, message=f"[{i}/{len(doc_files)}] Читаю: {doc.name}")
        try:
            text = _extract_doc(doc, tid)
            if text.strip():
                content_blocks.append((f"ДОКУМЕНТ: {doc.name}", text))
                _upd(tid, message=f"✓ {doc.name}: {len(text):,} символів")
            else:
                _upd(tid, message=f"[!] {doc.name}: порожній вміст")
        except Exception as exc:
            _upd(tid, message=f"[!] Помилка {doc.name}: {exc}")


# ── Claude analysis ───────────────────────────────────────────────────────

def _build_system_prompt(case: dict) -> str:
    if case.get("defendant") and case["defendant"] not in ("Прізвище Ім'я По-батькові", ""):
        case_block = f"""
ДАНІ СПРАВИ:
• Обвинувачений: {case.get('defendant', '—')}  ({case.get('dob', '—')})
• Справа: {case.get('case_number', '—')},  {case.get('article', '—')}
• Засідання: {case.get('hearing_date', '—')} о {case.get('hearing_time', '—')}
• Суддя: {case.get('judge', '—')},  {case.get('court', '—')}
"""
    else:
        case_block = ""

    return f"""Ти — досвідчений адвокат захисту з 15-річним стажем у судах України.
{case_block}
Тобі надані матеріали справи (аудіотранскрипції та/або документи).
Мова матеріалів — суміш українська + російська, аналізуй обидві.

Підготуй структурований пакет захисту у такому форматі:

=== 1. КЛАСИФІКАЦІЯ МАТЕРІАЛІВ ===
Для кожного аудіо/документа: СИЛЬНИЙ / ДОПОМІЖНИЙ / НЕЙТРАЛЬНИЙ / РИЗИКОВИЙ
(+ коротка причина + ключова цитата якщо є)

=== 2. ЗАГАЛЬНА КАРТИНА СПРАВИ ===
Що відбулось, хто сторони, суть обвинувачення.

=== 3. СИЛЬНІ ПОЗИЦІЇ ЗАХИСТУ ===
Конкретні факти, цитати, докази що допомагають обвинуваченому.

=== 4. СЛАБКІ МІСЦЯ / РИЗИКИ ===
Що може бути використано проти — будь чесним.

=== 5. СТРАТЕГІЯ ЗАХИСТУ ===
Конкретний план: що говорити, що подавати, якого порядку дотримуватись.

=== 6. КЛЮЧОВІ АРГУМЕНТИ (для виголошення в залі) ===
По пунктах, готові формулювання.

=== 7. КЛОПОТАННЯ ===
Які подавати, коли, щоб проситити.

=== 8. ПИТАННЯ ДО СВІДКІВ / ЗАЯВНИЦІ ===
Конкретні питання що підривають позицію обвинувачення.

=== 9. ШПАРГАЛКА ДО СУДУ ===
Одна сторінка найважливішого — для читання прямо в залі."""


def _analyze_case(
    tid: str,
    content_blocks: List[Tuple[str, str]],
    output_dir: Path,
) -> List[Path]:
    import anthropic, httpx

    try:
        from case_config import CASE
    except ImportError:
        try:
            from case_config_example import CASE  # type: ignore
        except ImportError:
            CASE = {}

    # Build combined text
    parts = []
    for label, text in content_blocks:
        parts.append(f"{'='*60}\n{label}\n{'='*60}\n{text}")
    full_text = "\n\n".join(parts)

    # Truncate to stay well within Claude's context
    if len(full_text) > 190_000:
        full_text = full_text[:190_000] + "\n\n[... скорочено через обсяг ...]"

    system_prompt = _build_system_prompt(CASE)

    _upd(tid, message=f"Надсилаю {len(full_text):,} символів у Claude...")

    client = anthropic.Anthropic(
        http_client=httpx.Client(verify=False, timeout=300.0)
    )

    result = ""
    with client.messages.stream(
        model="claude-sonnet-4-6",
        max_tokens=16000,
        system=system_prompt,
        messages=[{"role": "user", "content": f"МАТЕРІАЛИ ДЛЯ АНАЛІЗУ:\n\n{full_text}"}],
    ) as stream:
        for chunk in stream.text_stream:
            result += chunk

    _upd(tid, message=f"Claude відповів: {len(result):,} символів")

    output_dir.mkdir(parents=True, exist_ok=True)
    saved = []

    analysis_file = output_dir / "00_АНАЛІЗ_СПРАВИ.txt"
    analysis_file.write_text(result, encoding="utf-8")
    saved.append(analysis_file)
    _upd(tid, message=f"Збережено: {analysis_file.name}")

    # Extract cheat sheet
    m = re.search(r"={3,}\s*9[\.\s]+ШПАРГАЛКА[^\n]*\n(.*?)(?:={3,}|$)", result, re.DOTALL | re.IGNORECASE)
    cheat = m.group(1).strip() if m else ""
    if not cheat and len(result) > 1000:
        cheat = result[-3000:]
    if cheat:
        cheat_file = output_dir / "00_ШПАРГАЛКА_В_ЗАЛ_СУДУ.txt"
        cheat_file.write_text(cheat, encoding="utf-8")
        saved.append(cheat_file)
        _upd(tid, message=f"Збережено: {cheat_file.name}")

    return saved


# ── Pipeline ──────────────────────────────────────────────────────────────

def _pipeline(tid: str):
    job_dir    = JOBS_DIR / tid
    upload_dir = job_dir / "uploads"
    output_dir = job_dir / "output"

    all_files = sorted(upload_dir.glob("*")) if upload_dir.exists() else []
    audio_files = [f for f in all_files if f.suffix.lower() in AUDIO_EXTS]
    doc_files   = [f for f in all_files if f.suffix.lower() in DOC_EXTS]

    if not audio_files and not doc_files:
        _upd(tid, stage="error", error="Немає файлів для обробки", progress=100,
             message="[!] Завантаж аудіо або документи (.pdf, .docx, .txt)")
        return

    content_blocks: List[Tuple[str, str]] = []

    # ── Step 1: Transcribe audio (5 → 45%) ───────────────────────────────
    if audio_files:
        _upd(tid, stage="transcribing",
             label=f"Транскрипція аудіо ({len(audio_files)} файл(ів))", progress=5,
             message=f"Аудіофайлів: {len(audio_files)}")
        _run_transcriptions(tid, audio_files, content_blocks, p_start=5, p_end=45)

    # ── Step 2: Extract docs (45 → 60%) ──────────────────────────────────
    if doc_files:
        _upd(tid, stage="extracting",
             label=f"Читання документів ({len(doc_files)})", progress=45,
             message=f"Документів: {len(doc_files)}")
        _run_extractions(tid, doc_files, content_blocks, p_start=45, p_end=60)

    if not content_blocks:
        _upd(tid, stage="error", progress=100,
             error="Не вдалося отримати текст з жодного файлу",
             message="[!] Перевір формат файлів або встанови залежності")
        return

    total_chars = sum(len(t) for _, t in content_blocks)
    _upd(tid, message=f"Зібрано {len(content_blocks)} блок(ів), {total_chars:,} символів")

    # ── Step 3: Claude analysis (60 → 99%) ───────────────────────────────
    _upd(tid, stage="analyzing", label="Аналіз справи (Claude)", progress=60,
         message="Надсилаю матеріали до Claude...")
    try:
        _analyze_case(tid, content_blocks, output_dir)
    except Exception as exc:
        _upd(tid, stage="error", error=str(exc), progress=100, message=f"[ERROR] {exc}")
        return

    # ── Done ──────────────────────────────────────────────────────────────
    files_out = []
    for f in sorted(output_dir.rglob("*.txt")):
        rel = str(f.relative_to(SCRIPT_DIR)).replace("\\", "/")
        files_out.append({"name": f.name, "path": rel, "kb": round(f.stat().st_size / 1024, 1)})

    _upd(tid, stage="completed", label="Готово!", progress=100,
         files=files_out, message="✅ Аналіз завершено!")


def start_pipeline(tid: str):
    Thread(target=_pipeline, args=(tid,), daemon=True).start()


def list_results() -> list:
    results = []
    if not JOBS_DIR.exists():
        return results
    for job in sorted(JOBS_DIR.iterdir(), reverse=True):
        out = job / "output"
        if out.exists():
            for f in sorted(out.rglob("*.txt")):
                rel = str(f.relative_to(SCRIPT_DIR)).replace("\\", "/")
                results.append({"name": f.name, "path": rel,
                                 "kb": round(f.stat().st_size / 1024, 1)})
    return results


def resolve_file(rel_path: str) -> Optional[Path]:
    p = (SCRIPT_DIR / rel_path.replace("/", os.sep)).resolve()
    try:
        p.relative_to(SCRIPT_DIR.resolve())
        return p if p.exists() else None
    except ValueError:
        return None
