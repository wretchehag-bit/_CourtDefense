"""
╔══════════════════════════════════════════════════════════════════════════╗
║         МАЙСТЕР-АГЕНТ ЗАХИСТУ — ПОВНИЙ АНАЛІЗ СПРАВИ                   ║
║              defense_master.py  |  ст. 173-2 КУпАП                     ║
╠══════════════════════════════════════════════════════════════════════════╣
║  PIPELINE:                                                               ║
║   1. SCAN    — знаходить усі .txt / .docx / результати advocate_agent   ║
║   2. INGEST  — читає кожен документ (python-docx для .docx)             ║
║   3. ANALYZE — Claude будує зведену картину справи                      ║
║   4. GENERATE— генерує фінальний пакет документів для суду              ║
╚══════════════════════════════════════════════════════════════════════════╝

Використання:
  python defense_master.py                        # повний режим
  python defense_master.py --scan-only            # тільки сканування (без API)
  python defense_master.py --root "D:\\12314234"  # кастомний шлях
"""

from __future__ import annotations

import os, sys, re, json, shutil, argparse, time, subprocess
from pathlib import Path
from datetime import datetime
from dataclasses import dataclass, field
from typing import Optional

if sys.stdout.encoding and sys.stdout.encoding.lower() not in ("utf-8", "utf-8-sig"):
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")

# ── deps ────────────────────────────────────────────────────────────────
try:
    import anthropic, httpx
except ImportError:
    print("[!] pip install anthropic httpx")
    sys.exit(1)

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[!] pip install python-docx  (для читання .docx)")

# ══════════════════════════════════════════════════════════════════════════
# КОНФІГУРАЦІЯ
# ══════════════════════════════════════════════════════════════════════════

OUTPUT_DIR = Path(__file__).parent / "ПАКЕТ_ЗАХИСТУ"

try:
    from case_config import CASE, ROOT_DATA_DIR
    DEFAULT_ROOT = ROOT_DATA_DIR
except ImportError:
    from case_config_example import CASE, ROOT_DATA_DIR  # type: ignore
    DEFAULT_ROOT = ROOT_DATA_DIR

# Файли що треба читати з кореня
PRIORITY_FILES = {
    # (glob-pattern, короткий опис, пріоритет 1=найважливіший)
    "all_pages_recognized.txt":       ("Повний текст справи 135 стор.", 1),
    "дело паша 135 старниц.txt":      ("Текст справи (дублікат)", 1),
    "*Прям*спростовують*.docx":       ("Спростування обвинувачень", 2),
    "*Опровержение*.docx":            ("Опровержение обвинений", 2),
    "*Вступне_слово*.docx":           ("Вступне слово для суду", 2),
    "*Комплексне_клопотання*.docx":   ("Комплексне клопотання", 2),
    "*Пакет_документів*.docx":        ("Пакет документів ВАД", 2),
    "*Проєкт_Судового*.docx":         ("Проєкт судового документа", 3),
    "*Договор*примирения*.docx":      ("Договір про примирення", 3),
    "*Мировое_соглашение*.docx":      ("Мирова угода", 3),
    "*Аудиозапись*2*.docx":           ("Аудіозапис №2 розшифровка", 3),
    "*звернення_до_суду*.docx":       ("Звернення до суду ч1", 3),
    "*дітей*.docx":                   ("Документи про опіку", 3),
}

# Папки транскрипцій аудіо (вже оброблені advocate_agent)
TRANSCRIPTS_DIRS = [
    r"нарезки\Отсортированные_данные_для_суда",
    r"ВІДІБРАНІ_ДОКАЗИ",   # результати advocate_agent якщо є поруч
]

# Ліміти контексту
MAX_DOC_CHARS   = 12_000   # максимум символів з одного документа
MAX_TOTAL_CHARS = 180_000  # загальний ліміт контексту (~45k токенів)

# ══════════════════════════════════════════════════════════════════════════
# СТРУКТУРИ ДАНИХ
# ══════════════════════════════════════════════════════════════════════════

@dataclass
class SourceDocument:
    path: Path
    label: str          # короткий опис
    priority: int       # 1=найважливіший
    doc_type: str       # "case_file" | "defense_doc" | "transcript" | "audio_analysis"
    text: str = ""
    char_count: int = 0
    read_error: Optional[str] = None

    def __post_init__(self):
        self.char_count = len(self.text)


@dataclass
class CaseAnalysis:
    """Результат аналізу всієї справи агентом."""
    case_overview: str = ""          # суть справи в 10 реченнях
    charges: list[str] = field(default_factory=list)          # що інкримінують
    prosecution_evidence: list[str] = field(default_factory=list)  # докази обвинувачення
    defense_strong_points: list[str] = field(default_factory=list) # сильні місця захисту
    contradictions: list[str] = field(default_factory=list)        # суперечності в матеріалах
    procedural_violations: list[str] = field(default_factory=list) # процесуальні порушення
    client_rebuttals: list[str] = field(default_factory=list)      # що клієнт спростовує
    recommended_documents: list[dict] = field(default_factory=list)# які доки готувати
    top_audio_evidence: list[str] = field(default_factory=list)    # топ аудіодоказів
    hearing_strategy: str = ""       # тактика на засіданні
    risk_assessment: str = ""        # оцінка ризиків


# ══════════════════════════════════════════════════════════════════════════
# ЧИТАННЯ ДОКУМЕНТІВ
# ══════════════════════════════════════════════════════════════════════════

def read_docx(path: Path) -> str:
    """Читає .docx через python-docx."""
    if not DOCX_AVAILABLE:
        return f"[python-docx не встановлено: pip install python-docx]"
    try:
        doc = DocxDocument(str(path))
        paragraphs = []
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                paragraphs.append(t)
        # Також читаємо таблиці
        for table in doc.tables:
            for row in table.rows:
                row_texts = [c.text.strip() for c in row.cells if c.text.strip()]
                if row_texts:
                    paragraphs.append(" | ".join(row_texts))
        return "\n".join(paragraphs)
    except Exception as e:
        return f"[Помилка читання DOCX: {e}]"


def read_txt(path: Path) -> str:
    """Читає .txt з автодетекцією кодування."""
    for enc in ("utf-8", "utf-8-sig", "cp1251", "cp1252", "latin1"):
        try:
            return path.read_text(encoding=enc, errors="strict")
        except (UnicodeDecodeError, LookupError):
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def truncate(text: str, max_chars: int, label: str = "") -> str:
    """Обрізає текст з позначкою."""
    if len(text) <= max_chars:
        return text
    half = max_chars // 2
    return (text[:half] +
            f"\n\n... [СКОРОЧЕНО: {label}, показано {max_chars:,}/{len(text):,} символів] ...\n\n" +
            text[-half:])


# ══════════════════════════════════════════════════════════════════════════
# СКАНУВАННЯ
# ══════════════════════════════════════════════════════════════════════════

def scan_all_documents(root: Path) -> list[SourceDocument]:
    """Знаходить і читає всі релевантні документи."""
    docs: list[SourceDocument] = []
    seen_paths: set[Path] = set()

    def add_doc(path: Path, label: str, priority: int, doc_type: str):
        if path in seen_paths or not path.exists():
            return
        seen_paths.add(path)

        if path.suffix.lower() == ".docx":
            text = read_docx(path)
        elif path.suffix.lower() == ".txt":
            text = read_txt(path)
        else:
            return

        text = text.strip()
        if not text or len(text) < 50:
            return

        docs.append(SourceDocument(
            path=path,
            label=label,
            priority=priority,
            doc_type=doc_type,
            text=truncate(text, MAX_DOC_CHARS, label),
            char_count=len(text),
        ))

    # ── 1. Пріоритетні файли з кореня ──────────────────────────────────
    print("  [SCAN] Пріоритетні документи справи...")
    ilovepdf_dir = root / "ilovepdf_extracted-pages (1)"

    # Текст справи (135 стор.) — читаємо повністю, це найважливіше
    for name in ["all_pages_recognized.txt", "дело паша 135 старниц.txt"]:
        for search_dir in [ilovepdf_dir, root]:
            p = search_dir / name
            if p.exists():
                text = read_txt(p).strip()
                if text:
                    seen_paths.add(p)
                    docs.append(SourceDocument(
                        path=p, label=f"СПРАВА 135 стор: {name}",
                        priority=1, doc_type="case_file",
                        text=truncate(text, 40_000, "справа 135 стор"),
                        char_count=len(text),
                    ))
                    print(f"    ✓ {name} ({len(text):,} символів)")
                    break

    # ── 2. Всі .docx в корені ──────────────────────────────────────────
    print("  [SCAN] Документи захисту (.docx)...")
    for docx_path in sorted(root.glob("*.docx")):
        # Визначаємо пріоритет за назвою
        name_lower = docx_path.name.lower()
        if any(k in name_lower for k in ["спростовують", "опровержение", "вступне", "комплексне"]):
            pri, dtype = 2, "defense_doc"
        elif any(k in name_lower for k in ["пакет", "проєкт", "клопотання"]):
            pri, dtype = 2, "defense_doc"
        elif any(k in name_lower for k in ["договор", "мировое", "примирение", "угода"]):
            pri, dtype = 3, "defense_doc"
        else:
            pri, dtype = 4, "defense_doc"
        add_doc(docx_path, docx_path.name, pri, dtype)
        print(f"    ✓ {docx_path.name}")

    # ── 3. Транскрипції аудіо (відсортовані advocate_agent) ─────────────
    print("  [SCAN] Результати аналізу аудіо...")

    # Якщо є ВІДІБРАНІ_ДОКАЗИ поруч — читаємо тільки STRONG
    advocate_output = Path(__file__).parent / "ВІДІБРАНІ_ДОКАЗИ"
    if advocate_output.exists():
        strong_dir = advocate_output / "1_СИЛЬНІ_ДОКАЗИ"
        report_file = advocate_output / "ЗВІТ_ДЛЯ_СУДУ.txt"
        analysis_file = advocate_output / "АНАЛІЗ_ВСІХ_ФАЙЛІВ.txt"

        if report_file.exists():
            add_doc(report_file, "Звіт агента: топ аудіодокази", 1, "audio_analysis")
            print(f"    ✓ ЗВІТ_ДЛЯ_СУДУ.txt")
        if analysis_file.exists():
            add_doc(analysis_file, "Повний аналіз аудіо (177 файлів)", 2, "audio_analysis")
            print(f"    ✓ АНАЛІЗ_ВСІХ_ФАЙЛІВ.txt")
        if strong_dir.exists():
            # Читаємо тільки _АНАЛІЗ.txt файли (не самі транскрипції)
            for f in sorted(strong_dir.glob("*_АНАЛІЗ.txt"))[:15]:
                add_doc(f, f"Сильний доказ: {f.stem}", 2, "audio_analysis")
            print(f"    ✓ Сильні аудіодокази: {len(list(strong_dir.glob('*_АНАЛІЗ.txt')))} файлів")
    else:
        # Скануємо оригінальні транскрипції
        sorted_dir = root / "нарезки" / "Отсортированные_данные_для_суда"
        if sorted_dir.exists():
            count = 0
            for rec_dir in sorted(sorted_dir.iterdir()):
                if not rec_dir.is_dir():
                    continue
                tr_dir = rec_dir / "папка_с_транскрипцией"
                for txt in (tr_dir.glob("*.txt") if tr_dir.exists() else rec_dir.glob("*.txt")):
                    add_doc(txt, f"Транскрипція: {rec_dir.name}", 3, "transcript")
                    count += 1
                    if count >= 20:  # ліміт для контексту
                        break
                if count >= 20:
                    break
            print(f"    ✓ Транскрипцій (до 20): {count}")

    # ── 4. Відправки (ZIP розпаковані або папки) ────────────────────────
    vidpravka_dir = root / "відправка 4  чистосердечна розмова про все"
    if vidpravka_dir.exists():
        for f in sorted(vidpravka_dir.glob("*.txt"))[:3]:
            add_doc(f, f"Відправка 4: {f.name}", 3, "defense_doc")

    # Сортуємо за пріоритетом
    docs.sort(key=lambda d: (d.priority, d.path.name))
    return docs


# ══════════════════════════════════════════════════════════════════════════
# ПРОМПТИ
# ══════════════════════════════════════════════════════════════════════════

ANALYSIS_SYSTEM = """Ти — провідний адвокат-практик по справах ст. 173-2 КУпАП (домашнє насильство) та цивільних провадженнях про місце проживання дітей. 15 років досвіду в судах Київської області. Ти знаєш судову практику ЄСПЛ, ВСУ, практику Києво-Святошинського суду.

Твоя задача — провести повний адвокатський аналіз справи та підготувати конкретні рекомендації для захисту.

ПРАВОВА БАЗА:
- ст. 173-2 КУпАП: домашнє насильство — потрібно УМИСНЕ систематичне насильство
- ст. 251, 268, 280 КУпАП: вимоги до доказів та процесуальні права
- ст. 31 Конституції: право на запис розмов де ти учасник
- Закон України "Про запобігання та протидію домашньому насильству" 2017
- ЄСПЛ: справа Opuz v. Turkey — стандарти доказування
- Ухвала ВСУ № 5-158кс(15) — законність аудіозаписів

Аналізуй матеріали як практик: шукай процесуальні порушення, суперечності, слабкі місця обвинувачення."""

ANALYSIS_USER = """Справа: {case_number}
Обвинувачений: {defendant} ({dob})
Засідання: {hearing_date} {hearing_time}, суддя {judge}, {court}
Стаття: {article}

Нижче — ВСІ матеріали справи. Проведи повний адвокатський аналіз.

{documents_block}

═══════════════════════════════════════
ЗАВДАННЯ — дай відповідь ТІЛЬКИ у форматі JSON (без markdown, без ```):
{{
  "case_overview": "суть справи в 5-8 реченнях — хто, що, коли, скільки протоколів",
  "charges": [
    "Протокол 1: [дата, суть, що інкримінують конкретно]",
    "Протокол 2: ...",
    "..."
  ],
  "prosecution_evidence": [
    "доказ обвинувачення 1 з посиланням на документ",
    "..."
  ],
  "defense_strong_points": [
    "сильний аргумент захисту 1 — з конкретним посиланням на матеріали",
    "..."
  ],
  "contradictions": [
    "суперечність 1: [що каже обвинувачення vs що є в матеріалах]",
    "..."
  ],
  "procedural_violations": [
    "порушення 1: [яка норма порушена, де в матеріалах]",
    "..."
  ],
  "client_rebuttals": [
    "клієнт спростовує: [конкретний епізод] → [його аргумент з документів]",
    "..."
  ],
  "recommended_documents": [
    {{"title": "назва документу", "type": "клопотання|пояснення|заява", "priority": 1, "reason": "навіщо потрібен", "key_points": ["пункт 1", "пункт 2"]}},
    ...
  ],
  "top_audio_evidence": [
    "назва запису — чому сильний доказ і як використати в суді",
    "..."
  ],
  "hearing_strategy": "тактика на засіданні 01.06.2026 — що говорити першим, як реагувати на заявницю, ключові моменти",
  "risk_assessment": "оцінка ризиків по кожному протоколу і загальний прогноз"
}}"""


DOCUMENT_GENERATOR_SYSTEM = """Ти — досвідчений адвокат, що складає юридичні документи українською мовою відповідно до норм КУпАП, КАС України та практики українських судів.

Складай документи:
- Юридично грамотно, з посиланнями на конкретні норми
- Без "води" та загальних фраз — тільки конкретика
- З цитатами з матеріалів справи там де це підсилює позицію
- У форматі готовому для подання в суд"""


# ══════════════════════════════════════════════════════════════════════════
# АНАЛІЗ СПРАВИ
# ══════════════════════════════════════════════════════════════════════════

def build_documents_block(docs: list[SourceDocument], max_total: int) -> tuple[str, int]:
    """Будує блок документів для промпту з урахуванням ліміту."""
    parts = []
    total = 0
    included = 0

    for doc in docs:
        if total + doc.char_count > max_total:
            # Вставляємо скорочену версію
            remaining = max_total - total
            if remaining > 500:
                text = doc.text[:remaining] + f"\n[... скорочено через ліміт контексту ...]"
                parts.append(f"\n{'═'*60}\n[{doc.doc_type.upper()}] {doc.label}\n{'─'*60}\n{text}")
                total += remaining
            break
        parts.append(f"\n{'═'*60}\n[{doc.doc_type.upper()}] {doc.label}\n{'─'*60}\n{doc.text}")
        total += doc.char_count
        included += 1

    return "\n".join(parts), included


def analyze_case(
    docs: list[SourceDocument],
    client: anthropic.Anthropic,
) -> CaseAnalysis:
    """Аналізує всі матеріали справи одним викликом Claude."""

    documents_block, included = build_documents_block(docs, MAX_TOTAL_CHARS)
    total_chars = sum(d.char_count for d in docs[:included])

    print(f"  Відправляємо {included}/{len(docs)} документів")
    print(f"  Загальний контекст: {total_chars:,} символів (~{total_chars//4:,} токенів)")

    user_msg = ANALYSIS_USER.format(**CASE, documents_block=documents_block)

    print("  Аналіз Claude Opus... (може зайняти 1-2 хв)")
    start = time.time()

    response = client.messages.create(
        model="claude-opus-4-8",
        max_tokens=8000,
        system=ANALYSIS_SYSTEM,
        messages=[{"role": "user", "content": user_msg}],
    )

    elapsed = time.time() - start
    print(f"  Аналіз завершено за {elapsed:.0f}с")

    raw = response.content[0].text.strip()
    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)

    # Зберігаємо сирий відповідь для дебагу
    raw_path = OUTPUT_DIR / "_raw_analysis.json"
    raw_path.parent.mkdir(parents=True, exist_ok=True)
    raw_path.write_text(raw, encoding="utf-8")

    try:
        data = json.loads(raw)
    except json.JSONDecodeError as e:
        print(f"  [!] JSON помилка: {e}. Дивись _raw_analysis.json")
        # Fallback — повертаємо що є
        return CaseAnalysis(case_overview=raw[:2000])

    return CaseAnalysis(
        case_overview=data.get("case_overview", ""),
        charges=data.get("charges", []),
        prosecution_evidence=data.get("prosecution_evidence", []),
        defense_strong_points=data.get("defense_strong_points", []),
        contradictions=data.get("contradictions", []),
        procedural_violations=data.get("procedural_violations", []),
        client_rebuttals=data.get("client_rebuttals", []),
        recommended_documents=data.get("recommended_documents", []),
        top_audio_evidence=data.get("top_audio_evidence", []),
        hearing_strategy=data.get("hearing_strategy", ""),
        risk_assessment=data.get("risk_assessment", ""),
    )


# ══════════════════════════════════════════════════════════════════════════
# ГЕНЕРАЦІЯ ДОКУМЕНТІВ
# ══════════════════════════════════════════════════════════════════════════

def generate_document(
    title: str,
    doc_type: str,
    analysis: CaseAnalysis,
    extra_context: str,
    client: anthropic.Anthropic,
) -> str:
    """Генерує один судовий документ."""

    context_block = f"""
АНАЛІЗ СПРАВИ (підсумок агента):
Суть справи: {analysis.case_overview}

Сильні аргументи захисту:
{chr(10).join(f'- {a}' for a in analysis.defense_strong_points[:8])}

Суперечності в матеріалах:
{chr(10).join(f'- {c}' for c in analysis.contradictions[:5])}

Процесуальні порушення:
{chr(10).join(f'- {v}' for v in analysis.procedural_violations[:5])}

Спростування клієнта:
{chr(10).join(f'- {r}' for r in analysis.client_rebuttals[:5])}

{extra_context}
"""

    prompt = f"""Справа: {CASE['case_number']} | {CASE['article']}
Обвинувачений: {CASE['defendant']} ({CASE['dob']})
Засідання: {CASE['hearing_date']} {CASE['hearing_time']}
Суд: {CASE['court']}, суддя {CASE['judge']}

{context_block}

Склади документ: "{title}" (тип: {doc_type})

Вимоги:
- Мова: УКРАЇНСЬКА
- Формат: готовий до подання в суд
- Конкретні посилання на норми закону
- Без загальних фраз — тільки факти і юридичні аргументи
- Обсяг: достатній для повного викладу позиції

Почни відразу з тексту документа (без пояснень від себе)."""

    response = client.messages.create(
        model="claude-opus-4-8",
        max_tokens=4000,
        system=DOCUMENT_GENERATOR_SYSTEM,
        messages=[{"role": "user", "content": prompt}],
    )
    return response.content[0].text.strip()


def generate_all_documents(analysis: CaseAnalysis, client: anthropic.Anthropic, output_dir: Path):
    """Генерує весь пакет документів на основі аналізу."""

    output_dir.mkdir(parents=True, exist_ok=True)

    # Визначаємо які документи генерувати
    # Беремо топ-6 з рекомендованих + обов'язкові
    docs_to_generate = [
        {
            "filename": "01_ЗВЕДЕНІ_ПИСЬМОВІ_ПОЯСНЕННЯ.txt",
            "title": "Письмові пояснення підзахисного",
            "type": "пояснення",
            "extra": f"Спростування по кожному протоколу:\n" +
                     "\n".join(f"- {r}" for r in analysis.client_rebuttals),
        },
        {
            "filename": "02_КЛОПОТАННЯ_АУДІОЗАПИСИ.txt",
            "title": "Клопотання про долучення аудіозаписів як доказів",
            "type": "клопотання",
            "extra": f"Топ аудіодокази:\n" +
                     "\n".join(f"- {a}" for a in analysis.top_audio_evidence[:10]),
        },
        {
            "filename": "03_КЛОПОТАННЯ_ПОВЕРНЕННЯ_ПРОТОКОЛУ.txt",
            "title": "Клопотання про повернення протоколів на доопрацювання",
            "type": "клопотання",
            "extra": f"Процесуальні порушення:\n" +
                     "\n".join(f"- {v}" for v in analysis.procedural_violations),
        },
        {
            "filename": "04_КЛОПОТАННЯ_СВІДКИ.txt",
            "title": "Клопотання про виклик свідків",
            "type": "клопотання",
            "extra": "Свідки що можуть підтвердити нормальний характер стосунків у сім'ї",
        },
        {
            "filename": "05_ВИСТУП_В_СУДІ.txt",
            "title": "Промова захисту (виступ в суді)",
            "type": "промова",
            "extra": f"Стратегія: {analysis.hearing_strategy}",
        },
        {
            "filename": "06_СПРОСТУВАННЯ_ОБВИНУВАЧЕНЬ.txt",
            "title": "Письмове спростування кожного епізоду обвинувачення",
            "type": "пояснення",
            "extra": f"Суперечності:\n" +
                     "\n".join(f"- {c}" for c in analysis.contradictions),
        },
    ]

    # Додаємо рекомендовані агентом документи (якщо є нові)
    existing_titles = {d["title"] for d in docs_to_generate}
    for i, rec in enumerate(analysis.recommended_documents[:4], 7):
        if rec.get("title") not in existing_titles:
            docs_to_generate.append({
                "filename": f"0{i}_РЕКОМЕНДОВАНИЙ_{rec.get('type','doc').upper()}.txt",
                "title": rec.get("title", f"Документ {i}"),
                "type": rec.get("type", "клопотання"),
                "extra": "Ключові пункти:\n" + "\n".join(f"- {p}" for p in rec.get("key_points", [])),
            })

    total = len(docs_to_generate)
    for i, doc_spec in enumerate(docs_to_generate, 1):
        print(f"  [{i}/{total}] Генерую: {doc_spec['title'][:50]}...")
        try:
            text = generate_document(
                title=doc_spec["title"],
                doc_type=doc_spec["type"],
                analysis=analysis,
                extra_context=doc_spec["extra"],
                client=client,
            )
            out_path = output_dir / doc_spec["filename"]
            out_path.write_text(text, encoding="utf-8")
            print(f"         ✓ {doc_spec['filename']} ({len(text):,} символів)")
        except Exception as e:
            print(f"         [!] Помилка: {e}")


# ══════════════════════════════════════════════════════════════════════════
# МАЙСТЕР-ЗВІТ
# ══════════════════════════════════════════════════════════════════════════

def save_master_report(analysis: CaseAnalysis, docs: list[SourceDocument], output_dir: Path):
    """Зберігає зведений аналітичний звіт."""

    lines = [
        "╔══════════════════════════════════════════════════════════════════════╗",
        "║                  МАЙСТЕР-АНАЛІЗ СПРАВИ                             ║",
        "╠══════════════════════════════════════════════════════════════════════╣",
        f"║  {CASE['case_number']:<68}║",
        f"║  Засідання: {CASE['hearing_date']} {CASE['hearing_time']}, {CASE['judge']:<43}║",
        "╚══════════════════════════════════════════════════════════════════════╝",
        f"\nДата аналізу: {datetime.now().strftime('%d.%m.%Y %H:%M')}",
        f"Проаналізовано документів: {len(docs)}",
        "",
        "═" * 70,
        "1. СУТЬ СПРАВИ",
        "═" * 70,
        analysis.case_overview,
        "",
        "═" * 70,
        f"2. ОБВИНУВАЧЕННЯ ({len(analysis.charges)} протоколів/епізодів)",
        "═" * 70,
    ]
    for i, c in enumerate(analysis.charges, 1):
        lines += [f"  {i}. {c}", ""]

    lines += [
        "═" * 70,
        f"3. СИЛЬНІ АРГУМЕНТИ ЗАХИСТУ ({len(analysis.defense_strong_points)})",
        "═" * 70,
    ]
    for i, a in enumerate(analysis.defense_strong_points, 1):
        lines += [f"  {i}. {a}", ""]

    lines += [
        "═" * 70,
        f"4. СУПЕРЕЧНОСТІ В МАТЕРІАЛАХ ({len(analysis.contradictions)})",
        "═" * 70,
    ]
    for i, c in enumerate(analysis.contradictions, 1):
        lines += [f"  {i}. {c}", ""]

    lines += [
        "═" * 70,
        f"5. ПРОЦЕСУАЛЬНІ ПОРУШЕННЯ ({len(analysis.procedural_violations)})",
        "═" * 70,
    ]
    for i, v in enumerate(analysis.procedural_violations, 1):
        lines += [f"  {i}. {v}", ""]

    lines += [
        "═" * 70,
        f"6. СПРОСТУВАННЯ КЛІЄНТА ({len(analysis.client_rebuttals)})",
        "═" * 70,
    ]
    for i, r in enumerate(analysis.client_rebuttals, 1):
        lines += [f"  {i}. {r}", ""]

    lines += [
        "═" * 70,
        f"7. ТОП АУДІОДОКАЗІВ ({len(analysis.top_audio_evidence)})",
        "═" * 70,
    ]
    for i, a in enumerate(analysis.top_audio_evidence, 1):
        lines += [f"  {i}. {a}", ""]

    lines += [
        "═" * 70,
        "8. ТАКТИКА НА ЗАСІДАННІ 01.06.2026",
        "═" * 70,
        analysis.hearing_strategy,
        "",
        "═" * 70,
        "9. ОЦІНКА РИЗИКІВ ТА ПРОГНОЗ",
        "═" * 70,
        analysis.risk_assessment,
        "",
        "═" * 70,
        "10. РЕКОМЕНДОВАНІ ДОКУМЕНТИ ДЛЯ ПІДГОТОВКИ",
        "═" * 70,
    ]
    for i, d in enumerate(analysis.recommended_documents, 1):
        lines += [
            f"  {i}. [{d.get('priority','?')}] {d.get('title','')} ({d.get('type','')})",
            f"     Навіщо: {d.get('reason','')}",
            "     Ключові пункти: " + "; ".join(d.get("key_points", [])),
            "",
        ]

    lines += [
        "═" * 70,
        "11. ДЖЕРЕЛА АНАЛІЗУ",
        "═" * 70,
    ]
    for d in docs:
        lines.append(f"  [{d.priority}] [{d.doc_type}] {d.label} ({d.char_count:,} символів)")

    path = output_dir / "00_МАЙСТЕР_АНАЛІЗ.txt"
    path.write_text("\n".join(lines), encoding="utf-8")
    print(f"  ✓ {path.name}")


def save_quick_reference(analysis: CaseAnalysis, output_dir: Path):
    """Коротка шпаргалка на 1 сторінку — взяти в зал суду."""
    lines = [
        "ШПАРГАЛКА ДЛЯ ЗАЛУ СУДУ — 01.06.2026 10:00",
        "=" * 55,
        f"Справа: {CASE['case_number']}",
        f"Суддя: {CASE['judge']}",
        "",
        "ЩО ГОВОРИТИ ПЕРШИМ:",
        "  → 'Клопочу про долучення аудіозаписів як доказів'",
        "  → 'Клопочу про повернення протоколів на доопрацювання'",
        "  → 'Вину не визнаю. Прошу дослідити всі докази.'",
        "",
        "КЛЮЧОВІ АРГУМЕНТИ:",
    ]
    for i, a in enumerate(analysis.defense_strong_points[:5], 1):
        lines.append(f"  {i}. {a[:100]}")
    lines += [
        "",
        "ПРОЦЕСУАЛЬНІ ПОРУШЕННЯ (назвати суддіI):",
    ]
    for v in analysis.procedural_violations[:3]:
        lines.append(f"  • {v[:100]}")
    lines += [
        "",
        "КАТЕГОРИЧНО НЕ МОЖНА:",
        "  ✗ Підвищувати голос",
        "  ✗ Перебивати суддю",
        "  ✗ Визнавати факти під тиском",
        "  ✗ Коментувати показання без дозволу",
        "",
        "ТАКТИКА:",
    ]
    # Перші 3 речення стратегії
    strategy_sentences = analysis.hearing_strategy.split('.')[:3]
    for s in strategy_sentences:
        if s.strip():
            lines.append(f"  → {s.strip()}.")
    lines += [
        "",
        "ПРОГНОЗ:",
        f"  {analysis.risk_assessment[:200]}",
    ]

    path = output_dir / "00_ШПАРГАЛКА_В_ЗАЛ_СУДУ.txt"
    path.write_text("\n".join(lines), encoding="utf-8")
    print(f"  ✓ {path.name}")


# ══════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(description="Майстер-агент захисту")
    parser.add_argument("--root",       type=str, default=str(DEFAULT_ROOT))
    parser.add_argument("--output",     type=str, default=str(OUTPUT_DIR))
    parser.add_argument("--scan-only",  action="store_true", help="Тільки сканування")
    parser.add_argument("--no-docs",    action="store_true", help="Аналіз без генерації документів")
    args = parser.parse_args()

    root_dir   = Path(args.root)
    output_dir = Path(args.output)
    output_dir.mkdir(parents=True, exist_ok=True)

    print("╔══════════════════════════════════════════════════════════════════╗")
    print("║         МАЙСТЕР-АГЕНТ ЗАХИСТУ — ПОВНИЙ АНАЛІЗ СПРАВИ          ║")
    print("╚══════════════════════════════════════════════════════════════════╝")
    print(f"  Справа: {CASE['case_number']}")
    print(f"  Засідання: {CASE['hearing_date']} {CASE['hearing_time']}")
    print(f"  Джерело: {root_dir}")
    print(f"  Результат: {output_dir}")
    print()

    # ── КРОК 1: Сканування ──────────────────────────────────────────────
    print("[1/4] СКАНУВАННЯ ДОКУМЕНТІВ...")
    docs = scan_all_documents(root_dir)
    total_chars = sum(d.char_count for d in docs)
    print(f"\n  Разом: {len(docs)} документів, {total_chars:,} символів")
    print(f"  Розподіл за типами:")
    for dtype in ["case_file", "defense_doc", "audio_analysis", "transcript"]:
        count = len([d for d in docs if d.doc_type == dtype])
        if count:
            print(f"    {dtype}: {count}")

    if args.scan_only:
        print("\n[--scan-only] Сканування завершено.")
        print("\nЗнайдені документи:")
        for d in docs:
            print(f"  [{d.priority}] {d.label} ({d.char_count:,} символів) — {d.path.name}")
        return

    if not docs:
        print("[!] Документів не знайдено.")
        sys.exit(1)

    # ── КРОК 2: Аналіз справи ───────────────────────────────────────────
    print("\n[2/4] АНАЛІЗ СПРАВИ (Claude Opus)...")
    client = anthropic.Anthropic(
        http_client=httpx.Client(verify=False, timeout=300.0)
    )

    analysis = analyze_case(docs, client)

    # ── КРОК 3: Збереження аналізу ──────────────────────────────────────
    print("\n[3/4] ЗБЕРЕЖЕННЯ АНАЛІЗУ...")
    save_master_report(analysis, docs, output_dir)
    save_quick_reference(analysis, output_dir)

    # ── КРОК 4: Генерація документів ────────────────────────────────────
    if not args.no_docs:
        print("\n[4/4] ГЕНЕРАЦІЯ ПАКЕТУ ДОКУМЕНТІВ...")
        docs_dir = output_dir / "ДОКУМЕНТИ"
        generate_all_documents(analysis, client, docs_dir)
    else:
        print("\n[--no-docs] Генерацію документів пропущено.")

    # ── Підсумок ────────────────────────────────────────────────────────
    all_generated = list(output_dir.rglob("*.txt"))
    print()
    print("╔══════════════════════════════════════════════════════════════════╗")
    print("║                        ГОТОВО!                                  ║")
    print("╠══════════════════════════════════════════════════════════════════╣")
    print(f"║  Проаналізовано:    {len(docs):3} документів                              ║")
    print(f"║  Аргументів:        {len(analysis.defense_strong_points):3} сильних                             ║")
    print(f"║  Порушень:          {len(analysis.procedural_violations):3} процесуальних                       ║")
    print(f"║  Документів:        {len(all_generated):3} згенеровано                             ║")
    print("╠══════════════════════════════════════════════════════════════════╣")
    print(f"║  📁 {str(output_dir)[:61]:<61}║")
    print("╚══════════════════════════════════════════════════════════════════╝")
    print()
    print("  ПЕРШОЧЕРГОВО ВІДКРИТИ:")
    print(f"  1. {output_dir / '00_ШПАРГАЛКА_В_ЗАЛ_СУДУ.txt'}")
    print(f"  2. {output_dir / '00_МАЙСТЕР_АНАЛІЗ.txt'}")
    print(f"  3. {output_dir / 'ДОКУМЕНТИ'}/")


if __name__ == "__main__":
    main()

